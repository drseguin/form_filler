# keyword_parser.py
import re
import json
import os
from pathlib import Path
import streamlit as st
import logging
from datetime import date, datetime
from excel_manager import excelManager
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from logs.logger_config import setup_logger
from llm_factory import get_llm_client

class keywordParser:
    """
    A parser class that processes various keywords and extracts data from Excel,
    handles user input, and processes templates and JSON data using '!' as a separator.
    """

    def __init__(self, excel_manager=None, excel_managers=None):
        """
        Initialize the keyword parser.

        Args:
            excel_manager: An instance of excelManager to use for Excel operations.
                           If None, a new instance will be created when needed.
            excel_managers: A dictionary mapping Excel filenames to excelManager instances.
                            This is used for the new format with Excel file specification.
        """
        self.logger = setup_logger('keyword_parser')
        self.excel_manager = excel_manager
        self.excel_managers = excel_managers or {}  # Dictionary of Excel managers by filename
        self.pattern = r'{{(.*?)}}'
        self.has_input_fields = False
        self.form_submitted = False
        self.word_document = None
        self.input_values = {}  # Store input values
        
        # Load configuration from config.json
        self.config = self._load_config()
        
        # Get paths from config
        self.templates_dir = Path(self.config["paths"]["templates"])
        self.json_dir = Path(self.config["paths"]["json"])
        self.ai_dir = Path(self.config["paths"]["ai"])
        
        # Ensure templates directory exists
        if not self.templates_dir.exists():
            self.templates_dir.mkdir(parents=True, exist_ok=True)
            self.logger.info(f"Created templates directory: {self.templates_dir}")
        
        # Ensure json directory exists
        if not self.json_dir.exists():
            self.json_dir.mkdir(parents=True, exist_ok=True)
            self.logger.info(f"Created json directory: {self.json_dir}")
        
        self.logger.info("Initialized keywordParser")
        
        # Log Excel managers if available
        if excel_managers:
            excel_files = list(excel_managers.keys())
            self.logger.info(f"Initialized with {len(excel_files)} Excel managers: {excel_files}")

    def _load_config(self) -> dict:
        """Load configuration from config.json file."""
        config_path = Path("config.json")
        try:
            with open(config_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            self.logger.error(f"Error loading config.json: {str(e)}")
            # Return default config if file can't be loaded
            return {
                "paths": {
                    "templates": "templates",
                    "json": "json",
                    "ai": "ai"
                },
                "spacy": {
                    "enabled": True,
                    "model": "en_core_web_trf",
                    "format_entities": True,
                    "paragraph_breaks": True,
                    "entity_styles": {
                        "PERSON": {"bold": True},
                        "ORG": {"bold": True, "underline": True},
                        "DATE": {"italic": True},
                        "MONEY": {"bold": True},
                        "PLACE": {"underline": True}
                    }
                }
            }

    def set_word_document(self, doc):
        """Set the word document for direct table insertion."""
        self.word_document = doc
        self.logger.info("Word document set for table insertion")

    def parse(self, input_string):
        """
        Parse input string and process any keywords found.

        Args:
            input_string: The string containing keywords to parse.

        Returns:
            Processed string with keywords replaced with their values.
            If a keyword is replaced with a table, returns a dictionary with
            'text' and 'table' keys.
        """
        if not input_string:
            return input_string

        # Find all keywords in the input string
        matches = list(re.finditer(self.pattern, input_string))

        # First handle all INPUT keywords
        input_keywords = []
        for match in matches:
            content = match.group(1)  # Content inside {{}}
            keyword = match.group(0)  # The full {{keyword}}
            parts = content.split("!", 1) # Use '!' as separator
            keyword_type = parts[0].strip().upper()

            if keyword_type == "INPUT":
                input_keywords.append((keyword, content))

        # If we have input fields, process them first using a form
        if input_keywords and not self.form_submitted:
            with st.form(key=f"input_form_{id(input_string)}"):
                st.subheader("Please provide input values:")

                # Create input fields and store their values
                temp_input_values = {}
                for keyword, content in input_keywords:
                    value = self._create_input_field(content)
                    temp_input_values[keyword] = value

                # Add submit button
                submit = st.form_submit_button("Submit")
                if submit:
                    self.input_values.update(temp_input_values) # Store values upon submission
                    self.form_submitted = True
                    st.rerun() # Rerun to process the rest of the keywords
                else:
                    # If not submitted, show message and don't process further yet
                    st.stop()


        # After processing inputs or if no inputs, process all keywords
        result = input_string
        table_to_insert = None
        table_keyword = None
        docx_template_to_insert = None
        docx_template_keyword = None

        # Add debug logging
        self.logger.info(f"Processing {len(matches)} keywords in string: '{input_string[:50]}...'")

        for match in matches:
            keyword = match.group(0)  # Full keyword with {{}}
            content = match.group(1)  # Content inside {{}}

            # Always check first if this exact keyword is in our input_values dictionary
            if keyword in self.input_values:
                self.logger.info(f"Found keyword '{keyword}' in input_values dictionary")
                replacement = self.input_values[keyword]
            else:
                self.logger.info(f"Processing keyword '{keyword}', content: '{content}'")
                replacement = self._process_keyword(content)

            # Check if we got a docx template path back
            if isinstance(replacement, dict) and "docx_template" in replacement:
                # For a Word document template, we want to remember it but not do text replacement yet
                docx_template_to_insert = replacement["docx_template"]
                docx_template_keyword = keyword
                self.logger.info(f"Found template to insert from: {docx_template_to_insert}")
                # Don't do text replacement for this keyword yet
                continue
                
            # Check if we got a table object back
            elif isinstance(replacement, dict) and "table_object" in replacement:
                # For a table, we want to remember it but not do text replacement yet
                table_to_insert = replacement["table_object"]
                table_keyword = keyword
                self.logger.info(f"Found table to insert")
                # Don't do text replacement for this keyword yet
                continue
                
            # Regular text replacement
            # Ensure replacement is string, handle potential None values
            result = result.replace(keyword, str(replacement) if replacement is not None else "", 1)

        # Handle template insertion with priority over table
        if docx_template_to_insert and result.strip() == input_string.strip():
            # If the only content was the template keyword, return a special object
            self.logger.info(f"Returning template-only replacement from {docx_template_to_insert}")
            return {"text": "", "docx_template": docx_template_to_insert, "keyword": docx_template_keyword}
        elif docx_template_to_insert:
            # If there was a template keyword and other text, still return both
            # Replace the template keyword with an empty string in the result
            result = result.replace(docx_template_keyword, "", 1)
            self.logger.info(f"Returning mixed template and text replacement from {docx_template_to_insert}")
            return {"text": result, "docx_template": docx_template_to_insert, "keyword": docx_template_keyword}
        # Then handle table insertion if no template
        elif table_to_insert and result.strip() == input_string.strip():
            # If the only content was the table keyword, return a dict with both
            return {"text": "", "table": table_to_insert, "keyword": table_keyword}
        elif table_to_insert:
            # If there was a table keyword and other text, still return both
            # Replace the table keyword with an empty string in the result
            result = result.replace(table_keyword, "", 1)
            return {"text": result, "table": table_to_insert, "keyword": table_keyword}
        else:
            # Just regular text replacements happened
            return result

    def _create_input_field(self, content):
        """
        Create an appropriate input field based on the INPUT keyword using '!' separator.

        Args:
            content: The content inside the {{ }} brackets.

        Returns:
            The value from the input field.
        """
        if not content:
            return "[Invalid input reference]"

        # Split the content into tokens using '!'
        tokens = content.split("!")
        if len(tokens) < 2:
            return "[Invalid INPUT format]"

        # Get the keyword type (INPUT) and input type (text, area, date, select, check)
        keyword_type = tokens[0].strip().upper()
        input_type = tokens[1].strip().lower() if len(tokens) > 1 else ""

        # Check for valid INPUT keyword
        if keyword_type != "INPUT":
            return "[Invalid INPUT keyword]"
            
        # Create a consistent field key based on content
        field_key = f"input_field_{content}"

        # Handle text input - {{INPUT!text!label!value}}
        if input_type == "text":
            label = tokens[2] if len(tokens) > 2 else ""
            default_value = tokens[3] if len(tokens) > 3 else ""
            return st.text_input(
                label=label,
                value=default_value,
                label_visibility="visible",
                key=field_key
            )

        # Handle text area - {{INPUT!area!label!value!height}}
        elif input_type == "area":
            label = tokens[2] if len(tokens) > 2 else ""
            default_value = tokens[3] if len(tokens) > 3 else ""
            height_px = tokens[4] if len(tokens) > 4 else None

            # Convert height to integer if provided
            height = None
            if height_px:
                try:
                    height = int(height_px)
                except ValueError:
                    # If height is not a valid integer, ignore it
                    pass

            # Set height if provided, otherwise use default
            if height:
                return st.text_area(
                    label=label,
                    value=default_value,
                    height=height,
                    label_visibility="visible",
                     key=field_key
                )
            else:
                return st.text_area(
                    label=label,
                    value=default_value,
                    label_visibility="visible",
                     key=field_key
                )

        # Handle date input - {{INPUT!date!label!value!format}}
        elif input_type == "date":
            label = tokens[2] if len(tokens) > 2 else ""
            default_value_str = tokens[3] if len(tokens) > 3 else "today"
            date_format = tokens[4] if len(tokens) > 4 else "YYYY/MM/DD"

            # Handle "today" default value
            if default_value_str.lower() == "today":
                default_date = date.today()
            else:
                try:
                    # Try to parse the date based on the format
                    if date_format == "YYYY/MM/DD":
                        default_date = datetime.strptime(default_value_str, "%Y/%m/%d").date()
                    elif date_format == "DD/MM/YYYY":
                        default_date = datetime.strptime(default_value_str, "%d/%m/%Y").date()
                    elif date_format == "MM/DD/YYYY":
                        default_date = datetime.strptime(default_value_str, "%m/%d/%Y").date()
                    else:
                        # Default to ISO format if format is not recognized
                        default_date = datetime.strptime(default_value_str, "%Y-%m-%d").date()
                except ValueError:
                    default_date = date.today()

            date_value = st.date_input(
                label=label,
                value=default_date,
                label_visibility="visible",
                 key=field_key
            )

            # Return the date in the requested format
            if date_format == "YYYY/MM/DD":
                return date_value.strftime("%Y/%m/%d")
            elif date_format == "DD/MM/YYYY":
                return date_value.strftime("%d/%m/%Y")
            elif date_format == "MM/DD/YYYY":
                return date_value.strftime("%m/%d/%Y")
            else:
                return date_value.strftime("%Y/%m/%d")  # Default format

        # Handle select box - {{INPUT!select!label!options}}
        elif input_type == "select":
            label = tokens[2] if len(tokens) > 2 else ""
            options_str = tokens[3] if len(tokens) > 3 else ""

            # Parse options (comma-separated)
            options = [opt.strip() for opt in options_str.split(",")] if options_str else []

            if not options:
                return "[No options provided]"

            return st.selectbox(
                label=label,
                options=options,
                label_visibility="visible",
                 key=field_key
            )

        # Handle checkbox - {{INPUT!check!label!value}}
        elif input_type == "check":
            label = tokens[2] if len(tokens) > 2 else ""
            default_value_str = tokens[3].lower() if len(tokens) > 3 else "false"

            # Convert string value to boolean
            default_value = default_value_str == "true"

            return st.checkbox(
                label=label,
                value=default_value,
                label_visibility="visible",
                 key=field_key
            )

        # Default for unrecognized input types
        else:
            return f"[Unsupported input type: {input_type}]"

    def _process_keyword(self, content):
        """
        Process a single keyword content and return the corresponding value using '!' separator.

        Args:
            content: The content inside the {{ }} brackets.

        Returns:
            The processed value of the keyword.
        """
        parts = content.split("!", 1) # Use '!' as separator
        keyword_type = parts[0].strip().upper()

        # Process Excel data keywords
        if keyword_type == "XL":
            return self._process_excel_keyword(parts[1] if len(parts) > 1 else "")

        # Process user input keywords - these should already be handled in parse()
        elif keyword_type == "INPUT":
            # Fallback if not handled by the form (e.g., in tester_app without form)
             params = parts[1] if len(parts) > 1 else ""
             return self._process_input_keyword(params)


        # Process template keywords
        elif keyword_type == "TEMPLATE":
            return self._process_template_keyword(parts[1] if len(parts) > 1 else "")

        # Process JSON keywords
        elif keyword_type == "JSON":
            return self._process_json_keyword(parts[1] if len(parts) > 1 else "")
            
        # Process AI summary keywords
        elif keyword_type == "AI":
            return self._process_ai_keyword(parts[1] if len(parts) > 1 else "")

        # Unknown keyword type
        else:
            # If no known keyword type, treat it as a potential named range for XL
             return self._process_excel_keyword(f"RANGE!{content}")
            # return f"[Unknown keyword type: {keyword_type}]"

    def _process_input_keyword(self, params):
        """Process INPUT keywords directly if needed (fallback). Uses '!' separator."""
        input_parts = params.split("!") # Use '!' separator
        input_type = input_parts[0].lower() if input_parts else ""
        
        # First check if we have the full keyword in our input_values dictionary
        full_keyword = f"{{{{INPUT!{params}}}}}"
        if full_keyword in self.input_values:
            return self.input_values[full_keyword]
            
        # If not found with the INPUT! prefix, check if it's stored without the prefix
        # This handles cases where {{INPUT!text!Name!Joe}} might be stored as {{text!Name!Joe}} in input_values
        simple_keyword = f"{{{{{params}}}}}"
        if simple_keyword in self.input_values:
            return self.input_values[simple_keyword]

        # Fallback to default values if not found in input_values
        if input_type == "text" or input_type == "area":
            label = input_parts[1] if len(input_parts) > 1 else ""
            default_value = input_parts[2] if len(input_parts) > 2 else ""
            return default_value

        elif input_type == "date":
            # Use already imported datetime modules correctly
            today = date.today()
            return today.strftime("%Y/%m/%d")

        elif input_type == "select":
            options_str = input_parts[2] if len(input_parts) > 2 else ""
            options = [opt.strip() for opt in options_str.split(",")] if options_str else []
            return options[0] if options else ""

        elif input_type == "check":
            default_value_str = input_parts[2].lower() if len(input_parts) > 2 else "false"
            return default_value_str == "true"

        else:
            return params if params else "[Input value]"

    def _process_excel_keyword(self, content):
        """Process Excel-related keywords with new structure and '!' separator."""
        if not content:
            return "[Invalid Excel reference]"

        if not self.excel_manager:
            return "[Excel manager not initialized]"

        # Check if this is using the new format with Excel file specified 
        parts = content.split("!") 
        
        # Handle the case where the first part is an Excel file path
        excel_manager_to_use = self.excel_manager  # Default to the main Excel manager
        excel_file = None
        
        if parts[0].lower().endswith(('.xlsx', '.xls')):
            excel_file = parts[0]
            content_without_file = "!".join(parts[1:]) if len(parts) > 1 else ""
            
            # Check if we have this Excel file in our available managers dictionary
            if hasattr(self, 'excel_managers') and excel_file in self.excel_managers:
                excel_manager_to_use = self.excel_managers[excel_file]
                self.logger.info(f"Using Excel manager for file: {excel_file}")
                
                # Process the rest of the content using this manager
                return self._process_excel_content(content_without_file, excel_manager_to_use)
            else:
                # Try to load the file from the excel folder or current directory
                excel_dir = "excel"
                excel_path = os.path.join(excel_dir, excel_file)
                
                if os.path.exists(excel_path):
                    try:
                        # Create a temporary manager for this file
                        temp_manager = excelManager(excel_path)
                        self.logger.info(f"Created temporary Excel manager for file: {excel_file}")
                        
                        # Add to our managers dictionary if it exists
                        if hasattr(self, 'excel_managers'):
                            self.excel_managers[excel_file] = temp_manager
                        
                        # Process the content using this temp manager
                        result = self._process_excel_content(content_without_file, temp_manager)
                        
                        # Close the manager if it's not in our dictionary
                        if not hasattr(self, 'excel_managers'):
                            temp_manager.close()
                            
                        return result
                    except Exception as e:
                        self.logger.error(f"Error loading Excel file {excel_file}: {str(e)}", exc_info=True)
                        return f"[Error loading Excel file {excel_file}: {str(e)}]"
                elif os.path.exists(excel_file):
                    # Try from current directory
                    try:
                        # Create a temporary manager for this file
                        temp_manager = excelManager(excel_file)
                        self.logger.info(f"Created temporary Excel manager for file: {excel_file}")
                        
                        # Add to our managers dictionary if it exists
                        if hasattr(self, 'excel_managers'):
                            self.excel_managers[excel_file] = temp_manager
                        
                        # Process the content using this temp manager
                        result = self._process_excel_content(content_without_file, temp_manager)
                        
                        # Close the manager if it's not in our dictionary
                        if not hasattr(self, 'excel_managers'):
                            temp_manager.close()
                            
                        return result
                    except Exception as e:
                        self.logger.error(f"Error loading Excel file {excel_file}: {str(e)}", exc_info=True)
                        return f"[Error loading Excel file {excel_file}: {str(e)}]"
                else:
                    return f"[Excel file not found: {excel_file}]"
        
        # If we reach here, it's using the old format or the Excel file specification is invalid
        return self._process_excel_content(content, excel_manager_to_use)
        
    def _process_excel_content(self, content, excel_manager):
        """Process Excel content using the provided Excel manager."""
        parts = content.split("!") # Use '!' separator
        if len(parts) < 2:
             # Attempt to handle old format or named range as RANGE
            if ':' in content: # Could be old range XL:Sheet!A1:B2 or XL:A1:B2
                 if '!' in content.split(':')[0]: # Old range with sheet XL:Sheet!A1:B2
                     sheet_ref, cell_range = content.split('!', 1)
                     return self._call_excel_method("RANGE", f"{sheet_ref}!{cell_range}", excel_manager)
                 else: # Old range without sheet XL:A1:B2
                     # Explicitly pass RANGE type for cell ranges without sheet name
                     return self._call_excel_method("RANGE", content, excel_manager)
            elif content.startswith(':'): # Old LAST format XL::A1 or XL::Sheet!A1
                 return self._call_excel_method("LAST", content[1:], excel_manager) # Remove leading ':'
            else: # Assume it's a named range or old cell format XL:A1 or XL:Sheet!A1
                if '!' in content: # Old cell with sheet XL:Sheet!A1
                     return self._call_excel_method("CELL", content, excel_manager)
                else: # Old cell without sheet XL:A1 or a named range
                    # Try as cell first, if error, treat as named range
                    try:
                       return self._call_excel_method("CELL", content, excel_manager)
                    except ValueError:
                       return self._call_excel_method("RANGE", content, excel_manager) # Treat as named range


        xl_type = parts[0].strip().upper()
        xl_params = "!".join(parts[1:]) # Rejoin remaining parts

        return self._call_excel_method(xl_type, xl_params, excel_manager)


    def _call_excel_method(self, xl_type, xl_params, excel_manager=None):
        """Helper function to call the appropriate excelManager method."""
        # Use provided excel_manager if given, otherwise use the default
        manager = excel_manager or self.excel_manager
        
        if not manager:
            return "[Excel manager not available]"
            
        available_sheets = manager.get_sheet_names()
        sheet_name_map = {sheet.lower(): sheet for sheet in available_sheets}

        try:
            # {{XL!CELL!A1}} or {{XL!CELL!Sheet2!B5}}
            if xl_type == "CELL":
                sheet_name, cell_ref = self._get_sheet_and_ref(xl_params, available_sheets[0], sheet_name_map)
                return manager.read_cell(sheet_name, cell_ref)

            # {{XL!LAST!A1}} or {{XL!LAST!Sheet2!B5}}
            # {{XL!LAST!sheet_name!A1!Title}}
            elif xl_type == "LAST":
                last_parts = xl_params.split("!")
                if len(last_parts) >= 3: # Title format: {{XL!LAST!sheet_name!A1!Title}}
                    sheet_name_ref = last_parts[0]
                    cell_ref = last_parts[1]
                    title = last_parts[2]
                    actual_sheet_name = sheet_name_map.get(sheet_name_ref.lower(), sheet_name_ref) # Allow direct sheet name or lookup
                    if actual_sheet_name not in available_sheets: return f"[Sheet not found: {actual_sheet_name}]"
                    return manager.read_title_total(actual_sheet_name, cell_ref, title)
                    return self.excel_manager.read_title_total(actual_sheet_name, cell_ref, title)
                else: # Basic LAST format: {{XL!LAST!A1}} or {{XL!LAST!Sheet2!B5}}
                    sheet_name, cell_ref = self._get_sheet_and_ref(xl_params, available_sheets[0], sheet_name_map)
                    return self.excel_manager.read_total(sheet_name, cell_ref)

            # {{XL!RANGE!Sales!C3:C7}} or {{XL!RANGE!named_range}}
            elif xl_type == "RANGE":
                sheet_name, range_ref = self._get_sheet_and_ref(xl_params, available_sheets[0], sheet_name_map)
                # Debug log
                if hasattr(self.excel_manager, 'logger'):
                    self.excel_manager.logger.info(f"Processing RANGE: sheet='{sheet_name}', range='{range_ref}'")
                
                # If range_ref doesn't contain ':' it's likely a named range or invalid
                if ':' not in range_ref:
                    # Attempt to read as named range (assuming excelManager handles it)
                    # Or handle named range lookup if excelManager doesn't
                     # For now, assume read_range might handle named ranges or error out.
                    if hasattr(self.excel_manager, 'logger'):
                        self.excel_manager.logger.info(f"Range appears to be a named range: {range_ref}")
                    
                try:
                    data = self.excel_manager.read_range(sheet_name, range_ref)
                    if hasattr(self.excel_manager, 'logger'):
                        self.excel_manager.logger.info(f"Successfully read data from range: {sheet_name}!{range_ref}")
                        
                    if self.word_document and data:
                        return self._create_word_table(data)
                    else:
                        return self._format_table(data)
                except Exception as e:
                    error_msg = f"Error reading range '{range_ref}' from sheet '{sheet_name}': {str(e)}"
                    if hasattr(self.excel_manager, 'logger'):
                        self.excel_manager.logger.error(error_msg)
                    return f"[{error_msg}]"

            # {{XL!COLUMN!Sales!A1,C1,E1}} or {{XL!COLUMN!Sales!Revenue,Expense,Profit!12}}
            elif xl_type == "COLUMN":
                col_parts = xl_params.split("!")
                if len(col_parts) < 2: return "[Invalid COLUMN format]"

                sheet_ref = col_parts[0]
                columns_input = col_parts[1].strip('"') # Cell refs or titles

                actual_sheet_name = sheet_name_map.get(sheet_ref.lower(), sheet_ref) # Allow direct sheet name or lookup
                if actual_sheet_name not in available_sheets: return f"[Sheet not found: {actual_sheet_name}]"

                start_row = None
                use_titles = False

                if len(col_parts) > 2: # Optional start row provided, implies using titles
                    try:
                        start_row = int(col_parts[2])
                        use_titles = True
                    except ValueError:
                        return "[Invalid start row for COLUMN]"
                else:
                    # Determine if using titles based on input format (heuristic: check for letters vs numbers)
                    # A more robust check might be needed, e.g., trying to parse as cell ref
                    if not any(char.isdigit() for char in columns_input.replace(',', '')):
                         use_titles = True
                         start_row = 1 # Default start row for titles if not specified
                    # else: use_titles = False (default)


                data = self.excel_manager.read_columns(actual_sheet_name, columns_input, use_titles=use_titles, start_row=start_row)

                if self.word_document and data:
                    return self._create_word_table(data)
                else:
                    return self._format_table(data)


            else:
                return f"[Unknown XL type: {xl_type}]"

        except Exception as e:
            self.excel_manager.logger.error(f"Error processing XL keyword '{content}': {str(e)}", exc_info=True)
            return f"[Error processing XL: {str(e)}]"

    def _get_sheet_and_ref(self, params, default_sheet, sheet_map):
        """Helper to extract sheet name and cell/range reference."""
        parts = params.split("!")
        if len(parts) > 1 and parts[0].strip("'").lower() in sheet_map:
            # Sheet name is explicitly provided
            sheet_name = sheet_map[parts[0].strip("'").lower()]
            reference = "!".join(parts[1:]) # Rejoin if ref itself contains '!'
            # Debug log
            if self.excel_manager and hasattr(self.excel_manager, 'logger'):
                self.excel_manager.logger.info(f"Sheet name provided: '{sheet_name}', reference: '{reference}'")
        else:
            # Use default sheet
            sheet_name = default_sheet
            reference = params
            # Debug log
            if self.excel_manager and hasattr(self.excel_manager, 'logger'):
                self.excel_manager.logger.info(f"Using default sheet: '{sheet_name}', reference: '{reference}'")
                
        return sheet_name, reference


    def _format_table(self, data):
        """
        Format the data as a formatted table for Word or text.
        """
        if self.word_document:
            return self._create_word_table(data)

        if not data or not isinstance(data, list) or not all(isinstance(row, list) for row in data):
             return str(data) # Return raw data if not a list of lists

        # Calculate column widths
        col_widths = [0] * (max(len(row) for row in data) if data else 0)
        for row in data:
            for i, cell in enumerate(row):
                cell_str = str(cell) if cell is not None else ""
                if i < len(col_widths):
                     col_widths[i] = max(col_widths[i], len(cell_str))


        # Create the table as a string
        result = []
        for row_index, row in enumerate(data):
            row_str = []
            for i, cell in enumerate(row):
                 cell_str = str(cell) if cell is not None else ""
                 # Basic alignment (numbers right, text left) - simplistic check
                 try:
                      # Attempt to convert to float, fails for non-numeric strings
                      float(cell_str.replace(',', '').replace('$', ''))
                      formatted = cell_str.rjust(col_widths[i])
                 except (ValueError, TypeError):
                      formatted = cell_str.ljust(col_widths[i])

                 if i < len(col_widths):
                     row_str.append(formatted)


            result.append(" | ".join(row_str))

            # Add a separator after the header row (if more than one row exists)
            if row_index == 0 and len(data) > 1:
                separator = ["-" * width for width in col_widths]
                result.append("-+-".join(separator))

        return "\n".join(result)

    def _create_word_table(self, data):
        """
        Create a visually appealing table and return it to be inserted at the keyword position.
        """
        if not data or not isinstance(data, list) or not all(isinstance(row, list) for row in data):
            return str(data) # Return raw data representation if not table format

        num_rows = len(data)
        num_cols = max(len(row) for row in data) if num_rows > 0 else 0
        if num_cols == 0: return "[Empty Table Data]"

        # Create the table but don't add it to the document yet
        table = None
        if self.word_document:
            # Create a table object to return
            table = self.word_document.add_table(rows=num_rows, cols=num_cols)
            # Remove the table from the document (we'll insert it at the proper position later)
            table._element.getparent().remove(table._element)
        else:
            # If no document (e.g., in tester app), fall back to text format
            return self._format_table(data)
        
        # Try to apply a table style in order of preference
        style_applied = False
        try:
            # First try Table Grid
            table.style = 'Table Grid'
            style_applied = True
        except Exception as e1:
            try:
                # Then try other common styles
                for style_name in ['LightGrid', 'Grid Table Light', 'TableNormal', 'Normal Table']:
                    try:
                        table.style = style_name
                        style_applied = True
                        break  # Success, exit the style loop
                    except Exception:
                        continue  # Try the next style
            except Exception as e2:
                # If all styles fail, log but continue with default styling
                if hasattr(self.excel_manager, 'logger'):
                    self.excel_manager.logger.warning(f"Could not apply any table style: {str(e1)}")
        
        # If no style was applied, manually add borders to all cells
        if not style_applied:
            try:
                # Make sure required imports are available
                if 'OxmlElement' not in globals() or 'qn' not in globals():
                    if hasattr(self.excel_manager, 'logger'):
                        self.excel_manager.logger.info("Importing required modules for manual borders")
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn
                
                # Function to add border
                def set_cell_border(cell, border_type="single", size=4):
                    # Set each edge of the cell
                    for edge in ['top', 'left', 'bottom', 'right']:
                        tcPr = cell._tc.get_or_add_tcPr()
                        tcBorders = tcPr.first_child_found_or_add(qn('w:tcBorders'))
                        border_elm = OxmlElement(f'w:{edge}')
                        border_elm.set(qn('w:val'), border_type)
                        border_elm.set(qn('w:sz'), str(size))
                        border_elm.set(qn('w:space'), '0')
                        border_elm.set(qn('w:color'), 'auto')
                        tcBorders.append(border_elm)
                
                # Apply borders to all cells
                for i in range(num_rows):
                    for j in range(num_cols):
                        set_cell_border(table.cell(i, j))
                
                if hasattr(self.excel_manager, 'logger'):
                    self.excel_manager.logger.info("Applied manual borders to table cells")
            except Exception as e:
                if hasattr(self.excel_manager, 'logger'):
                    self.excel_manager.logger.warning(f"Could not apply manual borders: {str(e)}")
                    
        # Set overall table properties for better appearance
        try:
            # Set the table to auto-fit contents
            table.autofit = True

            # Fill the table with data and apply formatting
            for i, row in enumerate(data):
                for j in range(num_cols): # Ensure all cells in the row are processed
                    cell_value = row[j] if j < len(row) else None # Handle ragged rows

                    # Format the cell value (handle None)
                    if cell_value is None:
                         cell_text = ""
                    elif isinstance(cell_value, (int, float)):
                        cell_text = f"{cell_value:,.2f}" # Format numbers nicely
                    else:
                        cell_text = str(cell_value)

                    cell = table.cell(i, j)
                    # Check if cell contains multiple paragraphs and clear extra ones
                    if len(cell.paragraphs) > 1:
                         for p in cell.paragraphs[1:]:
                              p.clear() # Remove extra default paragraphs
                    # Ensure there's at least one paragraph to write to
                    if not cell.paragraphs:
                         cell.add_paragraph()

                    run = cell.paragraphs[0].clear().add_run(cell_text) # Clear and add new run

                    # Apply consistent font size
                    run.font.size = Pt(10)

                    # Apply padding within cells (apply to paragraph format)
                    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(3)

                    # Format header row (first row)
                    if i == 0:
                        run.font.bold = True
                        # Add light gray shading to header row
                        try:
                            tcPr = cell._tc.get_or_add_tcPr()
                            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>')
                            tcPr.append(shading_elm)
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as e:
                            # If XML manipulation fails, still make header bold
                            if hasattr(self.excel_manager, 'logger'):
                                self.excel_manager.logger.warning(f"Could not apply header row styling: {str(e)}")
                    else:
                         # Right-align numbers for better readability
                         # More robust check for numbers including currency
                        is_numeric = False
                        try:
                             # Try converting after removing currency symbols and commas
                             float(str(cell_value).replace('$', '').replace(',', ''))
                             is_numeric = True
                        except (ValueError, TypeError):
                             is_numeric = False

                        if is_numeric:
                             cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        else:
                             cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Apply alternating row colors (excluding header)
            try:
                for i in range(1, num_rows):
                    if i % 2 != 0:  # Apply shading to odd rows (1, 3, 5...)
                        for j in range(num_cols):
                            cell = table.cell(i, j)
                            tcPr = cell._tc.get_or_add_tcPr()
                            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')  # Very light gray
                            tcPr.append(shading_elm)
            except Exception as e:
                # If alternating rows styling fails, continue with basic table
                if hasattr(self.excel_manager, 'logger'):
                    self.excel_manager.logger.warning(f"Could not apply alternating row colors: {str(e)}")

        except Exception as e:
            # If any formatting fails, log and continue with basic table
            if hasattr(self.excel_manager, 'logger'):
                self.excel_manager.logger.warning(f"Table formatting could not be applied: {str(e)}")

        # Return the table object to be inserted at the keyword position
        return {"table_object": table}


    def _parse_section_param(self, param_part):
        """
        Extract section names from a section parameter string.
        Supports both single section and section ranges.
        
        Format: section=SectionName
               section=StartSection:EndSection
               section=SectionName&title=true/false
        
        Returns:
            Dictionary with 'start', optionally 'end', and 'include_title' keys
        """
        if not param_part.startswith("section="):
            return None
        
        self.logger.info(f"Parsing section parameters: '{param_part}'")
        
        # Initialize result with default title value (true by default)
        result = {'include_title': True}
        
        # Split into parameter pairs
        param_pairs = param_part.split('&')
        self.logger.info(f"Parameter pairs: {param_pairs}")
        
        # First handle the section parameter (expected to be the first one)
        section_param = param_pairs[0]
        section_value = section_param[len("section="):].strip()
        self.logger.info(f"Extracted section value: '{section_value}'")
        
        # Check if it's a section range (contains :)
        if ":" in section_value:
            start_section, end_section = section_value.split(":", 1)
            result.update({
                'start': start_section.strip(),
                'end': end_section.strip()
            })
            self.logger.info(f"Section range: '{start_section.strip()}' to '{end_section.strip()}'")
        else:
            # Single section
            result.update({
                'start': section_value,
                'end': None
            })
            self.logger.info(f"Single section: '{section_value}'")
        
        # Process any additional parameters
        for pair in param_pairs[1:]:
            if '=' in pair:
                key, value = pair.split('=', 1)
                key = key.strip().lower()
                value = value.strip().lower()
                self.logger.info(f"Additional parameter: {key}={value}")
                
                # Handle title parameter
                if key == 'title':
                    result['include_title'] = value == 'true'
                    self.logger.info(f"Title parameter set to: {result['include_title']}")
        
        self.logger.info(f"Final section parameters: {result}")
        return result

    def _process_template_keyword(self, content):
        """Process template keywords using '!' separator."""
        if not content:
            return "[Invalid TEMPLATE reference]"

        try:
            # Split into filename and optional parameters using '!'
            parts = content.split("!")
            filename = parts[0].strip()
            
            self.logger.info(f"Processing TEMPLATE keyword with filename: '{filename}', parts: {parts}")

            # Handle library templates {{TEMPLATE!LIBRARY!template_name!version}}
            if filename.upper() == "LIBRARY":
                 if len(parts) > 1:
                     template_name = parts[1].strip() if len(parts) > 1 else ""
                     template_version = parts[2].strip() if len(parts) > 2 else "DEFAULT"
                     # Implement template library lookup here
                     return f"[Template Library: {template_name} (Version: {template_version})]"
                 return "[Invalid library template reference]"

            # Always look in the templates directory
            template_path = self.templates_dir / filename
            self.logger.info(f"Template path resolved to: {template_path}")

            # Check if file exists first
            if not template_path.exists():
                self.logger.warning(f"Template file not found: {template_path}")
                return f"[Template file not found: {template_path}]"
            
            self.logger.info(f"Template file exists: {template_path}")
            
            # Parse additional parameters
            param_part = "".join(parts[1:]) if len(parts) > 1 else ""
            has_section = param_part.startswith("section=")
            
            # For Word documents with section parameter
            if filename.lower().endswith('.docx') and has_section:
                # Get the section info (single section or section range)
                section_info = self._parse_section_param(param_part)
                start_section = section_info['start'] 
                end_section = section_info['end']
                
                if end_section:
                    self.logger.info(f"Processing section range from '{start_section}' to '{end_section}'")
                else:
                    self.logger.info(f"Processing single section '{start_section}'")
                
                try:
                    from docx import Document
                    doc = Document(template_path)
                    
                    # DEBUG: Log all paragraphs that look like headings or titles
                    self.logger.info(f"DEBUG: Looking for section '{start_section}' in document")
                    debug_headings = []
                    for i, para in enumerate(doc.paragraphs):
                        # Check if this is a heading-like paragraph
                        is_heading = para.style and "heading" in para.style.name.lower()
                        is_title = (para.text.strip() and 
                                   len(para.text.strip()) < 100 and 
                                   not para.text.strip().endswith('.') and
                                   not para.text.strip().endswith(','))
                        
                        if is_heading or is_title:
                            debug_headings.append(f"{i}: '{para.text.strip()}'")
                    
                    self.logger.info(f"DEBUG: All potential section headings found:")
                    for heading in debug_headings:
                        self.logger.info(f"  {heading}")
                    
                    # Find sections and extract content
                    found_start = False
                    found_end = False
                    section_start_index = -1
                    section_end_index = -1
                    
                    # Store all heading-like paragraphs for debugging
                    headings = []
                    
                    # First pass: find start and end sections by headings or standalone titles
                    for i, para in enumerate(doc.paragraphs):
                        # Check if this is a heading-like paragraph
                        is_heading = para.style and "heading" in para.style.name.lower()
                        is_title = (para.text.strip() and 
                                   len(para.text.strip()) < 100 and 
                                   not para.text.strip().endswith('.') and
                                   not para.text.strip().endswith(','))
                        
                        if is_heading or is_title:
                            heading_text = para.text.strip()
                            headings.append((i, heading_text))
                            
                            # For Debug - Add this line
                            self.logger.info(f"Comparing heading '{heading_text}' with section '{start_section}'")
                            
                            # Look for start section with exact matching as first priority
                            if not found_start and heading_text == start_section:
                                found_start = True
                                section_start_index = i + 1  # Start after this heading
                                self.logger.info(f"Found start section (exact match) at paragraph {i}: '{heading_text}'")
                            # Then try normalized comparison if exact match fails
                            elif not found_start:
                                # Normalize texts for comparison to handle apostrophes and special characters
                                norm_heading = self._normalize_text(heading_text)
                                norm_start_section = self._normalize_text(start_section)
                                
                                # Try exact match with normalized text
                                if norm_heading == norm_start_section:
                                    found_start = True
                                    section_start_index = i + 1  # Start after this heading
                                    self.logger.info(f"Found start section (normalized match) at paragraph {i}: '{heading_text}'")
                                # Then try if heading contains the section name
                                elif norm_start_section in norm_heading:
                                    found_start = True
                                    section_start_index = i + 1  # Start after this heading
                                    self.logger.info(f"Found start section (contains match) at paragraph {i}: '{heading_text}'")
                                # Finally try if section name contains the heading (for when user gives too much detail)
                                elif len(norm_start_section) > 5 and norm_heading in norm_start_section:
                                    found_start = True
                                    section_start_index = i + 1  # Start after this heading
                                    self.logger.info(f"Found start section (reverse match) at paragraph {i}: '{heading_text}'")
                            
                            # Look for end section if specified
                            elif found_start and end_section:
                                # First try exact match
                                if heading_text == end_section:
                                    found_end = True
                                    section_end_index = i  # End before this heading
                                    self.logger.info(f"Found end section (exact match) at paragraph {i}: '{heading_text}'")
                                else:
                                    # Try normalized comparison
                                    norm_heading = self._normalize_text(heading_text)
                                    norm_end_section = self._normalize_text(end_section)
                                    
                                    # Try exact match with normalized text
                                    if norm_heading == norm_end_section:
                                        found_end = True
                                        section_end_index = i  # End before this heading
                                        self.logger.info(f"Found end section (normalized match) at paragraph {i}: '{heading_text}'")
                                    # Then try if heading contains the section name
                                    elif norm_end_section in norm_heading:
                                        found_end = True
                                        section_end_index = i  # End before this heading
                                        self.logger.info(f"Found end section (contains match) at paragraph {i}: '{heading_text}'")
                            
                            # If we found the start and no specific end was requested,
                            # any subsequent heading ends the section
                            elif found_start and not end_section:
                                found_end = True
                                section_end_index = i
                                self.logger.info(f"Found next heading at paragraph {i}: '{heading_text}'")
                                break
                    
                    # If we didn't find start by heading match, try exact text match
                    if not found_start:
                        self.logger.info(f"Looking for exact text match for start section: '{start_section}'")
                        for i, para in enumerate(doc.paragraphs):
                            if start_section.lower() == para.text.strip().lower():
                                found_start = True
                                section_start_index = i + 1  # Start after this paragraph
                                self.logger.info(f"Found start section by exact match at paragraph {i}")
                                break
                    
                    # If we found start but not end, and end is specified, look for exact match
                    if found_start and not found_end and end_section:
                        self.logger.info(f"Looking for exact text match for end section: '{end_section}'")
                        for i in range(section_start_index, len(doc.paragraphs)):
                            if end_section.lower() == doc.paragraphs[i].text.strip().lower():
                                found_end = True
                                section_end_index = i
                                self.logger.info(f"Found end section by exact match at paragraph {i}")
                                break
                    
                    # If we found start but not end, use end of document
                    if found_start and not found_end:
                        section_end_index = len(doc.paragraphs)
                        self.logger.info(f"Using end of document as section end (paragraph {section_end_index})")
                    
                    # Log error if section not found
                    if not found_start:
                        self.logger.warning(f"Could not find section '{start_section}'")
                        if headings:
                            self.logger.info("Available headings:")
                            for idx, heading in headings:
                                self.logger.info(f"  Paragraph {idx}: '{heading}'")
                        return f"[Section '{start_section}' not found in {filename}]"
                    
                    # Extract the selected paragraphs
                    section_paragraphs = doc.paragraphs[section_start_index:section_end_index]
                    if not section_paragraphs:
                        self.logger.warning(f"No content found in section")
                        return f"[No content found in section]"
                        
                    self.logger.info(f"Extracted {len(section_paragraphs)} paragraphs")
                    
                    # Always create a document, regardless of paragraph count
                    from docx import Document
                    temp_doc = Document()
                    
                    # Add title with section name only if include_title is True
                    include_title = section_info.get('include_title', True)
                    self.logger.info(f"Include title parameter is set to: {include_title}")
                    
                    if include_title:
                        # Add title with section name
                        title_para = temp_doc.add_paragraph(start_section)
                        try:
                            title_para.style = 'Heading 1'
                        except:
                            # Manually style if needed
                            title_run = title_para.runs[0]
                            title_run.bold = True
                            title_run.font.size = Pt(16)
                    
                    # Copy all paragraphs with formatting
                    for para in section_paragraphs:
                        p = temp_doc.add_paragraph()
                        # Copy text and formatting
                        for run in para.runs:
                            r = p.add_run(run.text)
                            r.bold = run.bold
                            r.italic = run.italic
                            r.underline = run.underline
                            if run.font.size:
                                r.font.size = run.font.size
                            if run.font.name:
                                r.font.name = run.font.name
                            if run.font.color.rgb:
                                r.font.color.rgb = run.font.color.rgb
                        
                        # Copy paragraph formatting
                        try:
                            if para.style:
                                p.style = para.style.name
                            p.paragraph_format.alignment = para.paragraph_format.alignment
                            p.paragraph_format.left_indent = para.paragraph_format.left_indent
                            p.paragraph_format.right_indent = para.paragraph_format.right_indent
                            p.paragraph_format.space_before = para.paragraph_format.space_before
                            p.paragraph_format.space_after = para.paragraph_format.space_after
                        except:
                            pass
                    
                    # Save to temporary file
                    import tempfile
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                        section_path = tmp.name
                        temp_doc.save(section_path)
                    
                    self.logger.info(f"Created section document at {section_path}")
                    return {"docx_template": section_path}
                except ImportError:
                    self.logger.error("python-docx library not available")
                    return f"[Error: python-docx library not available]"
                except Exception as e:
                    self.logger.error(f"Error extracting section: {str(e)}", exc_info=True)
                    return f"[Error extracting section: {str(e)}]"
            # Handle Word documents with no parameters (include whole document)
            elif filename.lower().endswith('.docx') and not param_part:
                self.logger.info(f"Including entire document: {template_path}")
                # Return the template path to be inserted
                return {"docx_template": str(template_path)}
            else:
                # Unknown parameter
                self.logger.warning(f"Unknown parameter: {param_part}")
                return f"[Unknown parameter: {param_part}]"
        except Exception as e:
            self.logger.error(f"Error processing template: {str(e)}", exc_info=True)
            return f"[Error in TEMPLATE: {str(e)}]"


    def _process_json_keyword(self, content):
        """Process JSON keywords using '!' separator."""
        if not content:
            return "[Invalid JSON reference]"

        try:
             # Split into filename, path, and optional transformation using '!'
            parts = content.split("!")
            
            # Handle the case where the first part might be empty ({{JSON!!filename.json}})
            if parts[0].strip() == "" and len(parts) > 1:
                filename = parts[1].strip()
                json_path = parts[2].strip() if len(parts) > 2 else "$"  # Default to root path
                transform_type = parts[3].strip().upper() if len(parts) > 3 else None
            else:
                if len(parts) < 2: return "[Invalid JSON format: Filename and path required]"
                filename = parts[0].strip()
                json_path = parts[1].strip()
                transform_type = parts[2].strip().upper() if len(parts) > 2 else None

            # If no JSON path is specified, use root ($) as default to return the entire file
            if not json_path:
                json_path = "$"
                
            # Handle special case where path is $. (root with empty element)
            if json_path == "$." or json_path == "$.":
                json_path = "$"

            # Check if filename is from another reference
            if filename.startswith("{{") and filename.endswith("}}"):
                # Recursively parse the reference
                filename = self.parse(filename)

            # Check if file exists directly at the provided path
            json_file_path = Path(filename)
            if not json_file_path.exists():
                # If not, check in the json folder
                json_folder_path = self.json_dir / filename
                if json_folder_path.exists():
                    json_file_path = json_folder_path
                    self.logger.info(f"Found JSON file in json folder: {json_file_path}")
                else:
                    return f"[JSON file not found: {filename} (checked in current directory and json folder)]"

            # Read the JSON file
            with open(json_file_path, 'r', encoding='utf-8') as file:
                json_data = json.load(file)

            # If path is just $ (root), return the entire JSON data
            if json_path == "$":
                # If we have a transform, apply it to the whole JSON object
                if transform_type:
                    # Handle transformations for the entire JSON object
                    if transform_type.startswith("JOIN(") and transform_type.endswith(")"):
                        if isinstance(json_data, list):
                            delimiter = transform_type[5:-1]
                            return delimiter.join(str(x) for x in json_data if x is not None)
                        
                # Return the complete JSON data
                return json.dumps(json_data, indent=2)

            # Simplistic JSONPath implementation (needs a library for full support)
            if json_path.startswith("$."):
                path_parts = json_path[2:].split(".")
                current = json_data
                
                # Handle the case where the path is $.
                if not path_parts[0] and len(path_parts) == 1:
                    return json.dumps(json_data, indent=2)

                for part in path_parts:
                    # Skip empty parts (handles cases like $.property..property2)
                    if not part:
                        continue
                        
                    # Handle array indexing like array[0] or [*]
                    if "[" in part and part.endswith("]"):
                        key = part.split("[")[0]
                        index_str = part.split("[")[1][:-1]

                        # Handle accessing the array itself if key is empty
                        if key:
                            if key not in current: return f"[JSON key not found: {key}]"
                            current = current[key]
                            if not isinstance(current, list): return f"[JSON path error: {key} is not an array]"

                        # Handle index or wildcard
                        if index_str == '*':
                             # This simplistic implementation doesn't fully support complex [*] behavior
                             # It might just return the list itself or error if used mid-path incorrectly
                             # A proper JSONPath library is needed for full support
                             pass # 'current' remains the list for now
                        else:
                             try:
                                 index = int(index_str)
                                 if index >= len(current): return f"[JSON index out of bounds: {index}]"
                                 current = current[index]
                             except (ValueError, IndexError, TypeError):
                                 return f"[Invalid JSON array index: {index_str}]"
                    else:
                        # Handle dynamic property names using keywords
                        if part.startswith("{{") and part.endswith("}}"):
                            part = self.parse(part) # Recursively parse key

                        if not isinstance(current, dict) or part not in current:
                             return f"[JSON key not found: {part}]"
                        current = current[part]


                 # Check for transformations if specified as the third part
                if transform_type:
                    if transform_type == "SUM" and isinstance(current, list):
                        try:
                            # Attempt to sum, converting elements to float
                             return sum(float(str(x).replace(',','').replace('$','')) for x in current if x is not None)
                        except (ValueError, TypeError):
                            return f"[Cannot SUM non-numeric values in list]"

                    elif transform_type.startswith("JOIN(") and transform_type.endswith(")"):
                        delimiter = transform_type[5:-1]
                        if isinstance(current, list):
                            return delimiter.join(str(x) for x in current if x is not None)
                        return str(current) # Join on single item returns the item as string


                    elif transform_type.startswith("BOOL(") and transform_type.endswith(")"):
                         yes_no = transform_type[5:-1].split("/")
                         yes_text = yes_no[0] if len(yes_no) > 0 else "Yes"
                         no_text = yes_no[1] if len(yes_no) > 1 else "No"

                         # Handle boolean conversion robustly
                         bool_value = False
                         if isinstance(current, bool):
                              bool_value = current
                         elif isinstance(current, str):
                              bool_value = current.lower() in ['true', 'yes', '1', 'on']
                         elif isinstance(current, (int, float)):
                              bool_value = current != 0

                         return yes_text if bool_value else no_text

                # Return the final value if no transformation or if transformation failed
                return current


            else:
                 return f"[Invalid JSONPath (must start with $.): {json_path}]"


        except json.JSONDecodeError:
            return f"[Error decoding JSON file: {json_file_path}]"
        except Exception as e:
            self.excel_manager.logger.error(f"Error processing JSON keyword '{content}': {str(e)}", exc_info=True)
            return f"[Error in JSON: {str(e)}]"

    def _process_ai_keyword(self, content):
        """Process AI keywords using '!' separator."""
        if not content:
            return "[Invalid AI reference]"
            
        try:
            # Split content to get document, prompt, and parameters
            parts = content.split("!")
            if len(parts) < 2:
                return "[Invalid AI format: Source document and prompt required]"
            
            source_doc = parts[0].strip()
            prompt_ref = parts[1].strip()
            
            # Process parameters (words limit, section)
            params = {}
            if len(parts) > 2:
                param_part = parts[2]
                # Simple parameter parsing (key=value&key2=value2)
                for param in param_part.split("&"):
                    if "=" in param:
                        key, value = param.split("=", 1)
                        params[key.strip().lower()] = value.strip()
            
            # Get words limit with default to 100
            words_limit = 100
            if "words" in params:
                try:
                    words_limit = int(params["words"])
                except ValueError:
                    pass
            
            self.logger.info(f"AI processing with words limit: {words_limit}")
            
            # Process section parameter if present
            section_info = None
            if "section" in params:
                section_value = params["section"]
                if ":" in section_value:
                    # It's a section range
                    start_section, end_section = section_value.split(":", 1)
                    section_info = {
                        'start': start_section.strip(),
                        'end': end_section.strip()
                    }
                else:
                    # Single section
                    section_info = {
                        'start': section_value.strip(),
                        'end': None
                    }
            
            # Check if AI directory exists, create if not
            if not self.ai_dir.exists():
                self.ai_dir.mkdir(parents=True, exist_ok=True)
                self.logger.info(f"Created AI directory: {self.ai_dir}")
            
            # First check if source document exists
            source_path = Path(source_doc)
            if not source_path.exists():
                # If not found at specified path, check in ai folder
                ai_source_path = self.ai_dir / source_doc
                if ai_source_path.exists():
                    source_path = ai_source_path
                    self.logger.info(f"Found source document in AI folder: {source_path}")
                else:
                    return f"[Source document not found: {source_doc} (checked in current directory and ai folder)]"
            
            # Extract text from the document to summarize
            document_text = ""
            
            try:
                # Handle various document types
                if source_path.name.lower().endswith('.docx'):
                    from docx import Document
                    doc = Document(source_path)
                    
                    # If section info is provided, extract only that section
                    if section_info:
                        # Similar to template section extraction logic
                        start_section = section_info['start']
                        end_section = section_info['end']
                        
                        # Find the section(s) in the document
                        found_start = False
                        found_end = False
                        section_paragraphs = []
                        
                        # First pass: find sections by headings or standalone titles
                        for i, para in enumerate(doc.paragraphs):
                            # Skip empty paragraphs
                            if not para.text.strip():
                                continue
                                
                            # Check if this is a heading-like paragraph
                            is_heading = para.style and "heading" in para.style.name.lower()
                            is_title = (para.text.strip() and 
                                       len(para.text.strip()) < 100 and 
                                       not para.text.strip().endswith('.') and
                                       not para.text.strip().endswith(','))
                            
                            heading_text = para.text.strip()
                            
                            # Look for start section
                            if not found_start:
                                # Try exact match
                                if heading_text == start_section:
                                    found_start = True
                                    self.logger.info(f"Found start section at paragraph {i}: '{heading_text}'")
                                    section_paragraphs.append(para)
                                    continue
                                
                                # Try normalized comparison
                                norm_heading = self._normalize_text(heading_text)
                                norm_start_section = self._normalize_text(start_section)
                                
                                if (norm_heading == norm_start_section or 
                                    norm_start_section in norm_heading or 
                                    (len(norm_start_section) > 5 and norm_heading in norm_start_section)):
                                    found_start = True
                                    self.logger.info(f"Found start section (normalized match) at paragraph {i}: '{heading_text}'")
                                    section_paragraphs.append(para)
                                    continue
                            
                            # If we're in the section, collect paragraphs until end
                            elif not found_end:
                                # If end section is specified, check if we've reached it
                                if end_section:
                                    # Try exact match for end section
                                    if heading_text == end_section:
                                        found_end = True
                                        self.logger.info(f"Found end section at paragraph {i}: '{heading_text}'")
                                        break
                                    
                                    # Try normalized comparison for end section
                                    norm_heading = self._normalize_text(heading_text)
                                    norm_end_section = self._normalize_text(end_section)
                                    
                                    if (norm_heading == norm_end_section or 
                                        norm_end_section in norm_heading):
                                        found_end = True
                                        self.logger.info(f"Found end section (normalized match) at paragraph {i}: '{heading_text}'")
                                        break
                                
                                # If no end section or if this isn't the end, but we're in a new section
                                elif (not end_section and 
                                     (is_heading or 
                                      (is_title and heading_text != start_section and 
                                       self._normalize_text(heading_text) != self._normalize_text(start_section)))):
                                    found_end = True
                                    self.logger.info(f"Found next heading at paragraph {i}: '{heading_text}'")
                                    break
                                
                                # Add paragraph to section content
                                section_paragraphs.append(para)
                        
                        # Extract text from the section paragraphs
                        if found_start:
                            document_text = "\n".join([p.text for p in section_paragraphs if p.text.strip()])
                        else:
                            return f"[Section '{start_section}' not found in {source_doc}]"
                    else:
                        # Extract all text from the document
                        document_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                
                elif source_path.name.lower().endswith('.txt'):
                    # For text files, read directly
                    with open(source_path, 'r', encoding='utf-8') as file:
                        document_text = file.read()
                
                else:
                    return f"[Unsupported document type: {source_path}. Please use .docx or .txt files]"
            
            except Exception as e:
                self.logger.error(f"Error extracting text from document: {str(e)}", exc_info=True)
                return f"[Error extracting text: {str(e)}]"
            
            # If no text was extracted
            if not document_text.strip():
                return "[No text found to summarize]"
            
            # Get the prompt from file or use directly
            prompt_text = ""
            if prompt_ref.lower().endswith('.txt'):
                # Look for prompt file
                prompt_path = Path(prompt_ref)
                if not prompt_path.exists():
                    # Check in AI folder
                    ai_prompt_path = self.ai_dir / prompt_ref
                    if ai_prompt_path.exists():
                        prompt_path = ai_prompt_path
                    else:
                        return f"[Prompt file not found: {prompt_ref} (checked in current directory and ai folder)]"
                
                # Read prompt from file
                try:
                    with open(prompt_path, 'r', encoding='utf-8') as file:
                        prompt_text = file.read().strip()
                except Exception as e:
                    self.logger.error(f"Error reading prompt file: {str(e)}", exc_info=True)
                    return f"[Error reading prompt file: {str(e)}]"
            else:
                # Use the prompt text directly
                prompt_text = prompt_ref
            
            # Use the LLM client to generate the summary instead of calling OpenAI directly
            try:
                # Get the LLM client (OpenAI or Triton based on config)
                llm_client = get_llm_client()
                
                # Call the summarize method
                summary = llm_client.summarize(
                    text=document_text,
                    prompt=prompt_text,
                    max_words=words_limit,
                    temperature=0.5
                )
                
                # Check if spaCy formatting is enabled in the config
                spacy_config = self.config.get("spacy", {})
                spacy_enabled = spacy_config.get("enabled", False)
                
                # Apply spaCy formatting if enabled
                if spacy_enabled and summary:
                    return self._format_text_with_spacy(summary, spacy_config)
                
                return summary
                
            except Exception as e:
                self.logger.error(f"Error generating summary: {str(e)}", exc_info=True)
                return f"[Error generating summary: {str(e)}]"
                
        except Exception as e:
            self.logger.error(f"Error processing AI keyword: {str(e)}", exc_info=True)
            return f"[Error in AI: {str(e)}]"

    def _format_text_with_spacy(self, text, spacy_config):
        """
        Format text using spaCy NLP features for better readability.
        
        Args:
            text: Text to format
            spacy_config: spaCy configuration parameters
            
        Returns:
            Formatted text with improved structure and formatting
        """
        if not text:
            return text
            
        try:
            # Import spaCy
            import spacy
            from spacy.tokens import Doc
            
            # Get spaCy model name from config or use default
            model_name = spacy_config.get("model", "en_core_web_trf")
            
            # Load spaCy model
            try:
                nlp = spacy.load(model_name)
                self.logger.info(f"Loaded spaCy model: {model_name}")
            except OSError:
                # Download the model if it doesn't exist
                self.logger.warning(f"spaCy model {model_name} not found. Attempting to download...")
                
                # Display a notification to the user about the download
                download_message = st.warning(f"Downloading spaCy language model: {model_name}. This may take a few minutes depending on your internet connection and the model size...", icon="⏳")
                
                spacy.cli.download(model_name)
                nlp = spacy.load(model_name)
                
                # Update the message once download is complete
                download_message.success(f"Successfully downloaded and loaded spaCy model: {model_name}")
                self.logger.info(f"Downloaded and loaded spaCy model: {model_name}")
            
            # Process the text with spaCy
            doc = nlp(text)
            
            # Get formatting options from config
            format_entities = spacy_config.get("format_entities", True)
            paragraph_breaks = spacy_config.get("paragraph_breaks", True)
            entity_styles = spacy_config.get("entity_styles", {})
            
            # If we're not using the word document (for preview), return a simpler format
            if not self.word_document:
                # Just apply basic formatting to improve text output
                sentences = list(doc.sents)
                result = []
                
                # Group sentences into paragraphs based on content
                current_paragraph = []
                for sent in sentences:
                    sent_text = sent.text.strip()
                    if not sent_text:
                        continue
                        
                    # Start a new paragraph at meaningful breaks
                    if (paragraph_breaks and 
                        (len(current_paragraph) > 3 or  # Natural paragraph length
                         sent.text.startswith("•") or    # Bullet points
                         any(token.is_title for token in sent) or  # Likely a heading
                         len(sent) < 5)):  # Very short sentence likely a title
                        
                        if current_paragraph:
                            result.append(" ".join(current_paragraph))
                            current_paragraph = []
                    
                    current_paragraph.append(sent_text)
                
                # Add the last paragraph
                if current_paragraph:
                    result.append(" ".join(current_paragraph))
                    
                return "\n\n".join(result)
            
            # For Word document, we can create rich text formatting
            # We'll return formatted text that preserves entity information for the document
            # This could be extended to create a proper Word document object with formatting
            # if needed in the future
            
            # For now, we'll keep it simple and just improve paragraph structure
            sentences = list(doc.sents)
            result = []
            
            # Group sentences into paragraphs based on content
            current_paragraph = []
            for sent in sentences:
                sent_text = sent.text.strip()
                if not sent_text:
                    continue
                    
                # Start a new paragraph at meaningful breaks
                if (paragraph_breaks and 
                    (len(current_paragraph) > 3 or  # Natural paragraph length
                     sent.text.startswith("•") or    # Bullet points
                     any(token.is_title for token in sent) or  # Likely a heading
                     len(sent) < 5)):  # Very short sentence likely a title
                    
                    if current_paragraph:
                        result.append(" ".join(current_paragraph))
                        current_paragraph = []
                
                current_paragraph.append(sent_text)
            
            # Add the last paragraph
            if current_paragraph:
                result.append(" ".join(current_paragraph))
                
            return "\n\n".join(result)
            
        except ImportError:
            self.logger.warning("spaCy library not available. Install with 'pip install spacy'")
            return text
        except Exception as e:
            self.logger.error(f"Error formatting text with spaCy: {str(e)}", exc_info=True)
            return text

    def reset_form_state(self):
        """Reset the form submission state and clear cached values."""
        self.form_submitted = False
        self.input_values = {}

    def clear_input_cache(self):
        """Clear the cached user inputs stored in session state (for Streamlit apps)."""
        # Find all INPUT keys in session state derived from input fields and clear them
        keys_to_clear = [key for key in st.session_state.keys() if key.startswith("input_field_INPUT!")]
        for key in keys_to_clear:
            # Reset based on type - needs knowledge of type or safe reset
             if isinstance(st.session_state[key], bool):
                 st.session_state[key] = False # Default for checkbox
             elif isinstance(st.session_state[key], (int, float)):
                  st.session_state[key] = 0 # Default for number (if ever used)
             else:
                 st.session_state[key] = "" # Default for text/area/select/date


        # Also clear the parser's internal state
        self.reset_form_state()



    def get_excel_keyword_help(self):
        """
        Get help text explaining how to use Excel Keywords with '!' separator.

        Returns:
            A string with help information about available keywords.
        """
        help_text = f"""
# Excel Keywords
If Excel keywords `{{{{XL!...}}}}` are detected in the uploaded document, the system will look for and prompt for any missing Excel files in Step 2.

## Excel File Specification
You can now specify which Excel file to use for each Excel keyword:

### {{{{XL!excel_file.xlsx!CELL!Cell}}}}
The application will look for `excel_file.xlsx` in the current directory or the `excel` folder. If not found, it will prompt the user to upload it.

## Available Excel Keywords

### {{{{XL!excel_file.xlsx!CELL!`Cell`}}}}
Get a value from `Cell` (ex: A1) in the specified Excel file.

### {{{{XL!excel_file.xlsx!CELL!`Sheet`!`Cell`}}}}
Get a value from `Cell` (ex: A1) in `Sheet` of the specified Excel file.

### {{{{XL!excel_file.xlsx!LAST!`Cell`}}}}
Get the last non-empty value going down from `Cell` (ex: A1). Used for getting totals.

### {{{{XL!excel_file.xlsx!LAST!`Sheet`!`Cell`}}}}
Get the last non-empty value going down from `Cell` (ex: A1) in `Sheet`. Used for getting totals.

### {{{{XL!excel_file.xlsx!LAST!`Sheet`!`Cell`!`Title`}}}}
From `Cell` (ex: A1), on `Sheet` scan right until the `Title` is detected, then get the last non-empty value going down from the `Title` column. Used for getting totals.

### {{{{XL!excel_file.xlsx!RANGE!`Start Cell`:`End Cell`}}}}
Get values for the range starting at `Start Cell` (ex: A1) to the `End Cell` (ex: G13). A formated table is returned.

### {{{{XL!excel_file.xlsx!RANGE!`Sheet`!`Start Cell`:`End Cell`}}}}
Get values for the range starting at `Start Cell` (ex: A1) to the `End Cell` (ex: G13) in `Sheet`. A formated table is returned.

### {{{{XL!excel_file.xlsx!COLUMN!`Sheet`!`Cell 1`,`Cell 2`,`Cell 3`,...}}}}
Returns a formatted table with columns `Cell 1` (ex: A1),`Cell 2` (ex: C1),`Cell 3` (ex: F1)... from `Sheet` appended together. Row number must be the same for each. Example: {{{{XL!budget.xlsx!COLUMN!Support!C4,E4,J4}}}}.

### {{{{XL!excel_file.xlsx!COLUMN!`Sheet`!`Title 1`,`Title 2`,`Title 3`,...!`Row`}}}}
Returns a formatted table with columns with `Title 1` (ex: Item),`Title 2` (ex: HST),`Title 3` (ex: Total)... from `Sheet` appended together. The `Title` row is specified by `Row` (ex: 6). Example: {{{{XL!sales.xlsx!COLUMN!Distribution Plan!Unit,DHTC,Total!4}}}}.
"""
        return help_text 

    def get_input_keyword_help(self):
        """
        Get help text explaining how to use User Iniput Keywords with '!' separator.

        Returns:
            A string with help information about available keywords.
        """
        help_text = f"""
# User Input Keywords
If User Input keywords `{{{{INPUT!...}}}}` are detected in the uploaded document, the user will be prompt for input value(s) in Step 3.
### {{{{INPUT!TEXT!`label`!`default_value`}}}}
Prompt the user for a single-line text input with `label` and `default_value`.
### {{{{INPUT!AREA!`label`!`default_value`!`height`}}}}
Prompt the user for a multi-line text input with `label`, `default_value`, and `height (ex: 200)`.
### {{{{INPUT!DATE!`label`!`default_date`!`format`}}}}
Prompt the user for a date input with `label`, `default_date` (ex: 1990/01/01), and `format` (ex: YYYY/MM/DD).
### {{{{INPUT!SELECT!`label`!`option1`!`option2`!`option3`!`...`}}}}
Prompt the user for a dropdown selection with `label` and options `option1`, `option2`, `option3`, etc.
### {{{{INPUT!CHECK!`label`!`default_state`}}}}
Prompt the user for a checkbox input with `label` and `default_state` (ex: True).
"""
        return help_text 

    def get_template_keyword_help(self):
        """
        Get help text explaining how to use Template Keywords with '!' separator.

        Returns:
            A string with help information about available keywords.
        """
        help_text = f"""
# Template Keywords
If Template keywords `{{{{TEMPLATE!...}}}}` are detected in the uploaded document, the application will look for the specified template file(s) `(ex: filename.docx)` in the `{self.templates_dir}` folder.
### {{{{TEMPLATE!`filename.docx`}}}}
Inject the full document content.
### {{{{TEMPLATE!`filename.docx`!`section=heading`}}}}
Inject the content of the section named `heading` including the section heading.
### {{{{TEMPLATE!`filename.docx`!`section=heading`!`title=false`}}}}
Inject the content of the section named `heading` without the section heading if title is set to false.
### {{{{TEMPLATE!`filename.docx`!`section=heading_start:heading_end`}}}}
Inject the content of the sections from `heading_start` to `heading_end` including the section heading.
### {{{{TEMPLATE!`filename.docx`!`section=heading_start:heading_end&title=false`}}}}
Inject the content of the sections from `heading_start` to `heading_end` without the section heading if title is set to false.
"""
        return help_text 

    def get_json_keyword_help(self):
        """
        Get help text explaining how to use JSON Keywords with '!' separator.

        Returns:
            A string with help information about available keywords.
        """
        help_text = f"""
# JSON Keywords
If JSON keywords `{{{{JSON!...}}}}` are detected in the uploaded document, the application will look for the specified JSON file(s) `(ex: filename.json)` in the `{self.json_dir}` folder. The system will first look for the file at the specified path, and if not found, it will check in the '{self.json_dir}' directory.
### {{{{JSON!!`filename.json`}}}}
Inject the full JSON content. Note the double `!!` to indicate the full JSON content.
### {{{{JSON!!`filename.json`!`$.`}}}}
Alternative syntax that also injects the full JSON content (the path `$.` refers to the root).
### {{{{JSON!`filename.json`!`$.key`}}}}
Inject the content of the JSON path `key`. Example: {{{{JSON!!launch.json!\\$.configurations}}}}. Example: {{{{JSON!launch.json!$.configurations[0].name}}}}.
### {{{{JSON!`filename.json`!`$.key`!`SUM`}}}}
Sum the numeric values in the JSON path `key`. Example: {{{{JSON!sales.json!$.monthly_totals!SUM}}}}.   
### {{{{JSON!`filename.json`!`$.key`!`JOIN(, )`}}}}
Join the values in the JSON path `key` with a comma and space. Example: {{{{JSON!users.json!$.names!JOIN(, )}}}}.
### {{{{JSON!`filename.json`!`$.key`!`BOOL(Yes/No)`}}}}
Transform the boolean values in the JSON path `key` to custom text. Example: {{{{JSON!status.json!$.system_active!BOOL(Online/Offline)}}}}.
"""
        return help_text 

    def get_ai_keyword_help(self):
        """
        Get help text explaining how to use AI Keywords with '!' separator.

        Returns:
            A string with help information about available keywords.
        """
        help_text = f"""
# AI Keywords
If AI keywords `{{{{AI!...}}}}` are detected in the uploaded document, the application will look for the specified document(s) in the `{self.ai_dir}` folder or at the specified path. If the '{self.ai_dir}' folder does not exist, it will be created automatically.

### {{{{AI!`source-doc.docx`!`prompt_file.txt`!`words=100`}}}}
Summarize the entire document located at '{self.ai_dir}/source-doc.docx'. The summary will be limited to 100 words or less. The prompt for the summary can be found in '{self.ai_dir}/prompt_file.txt'. If the prompt file does not have a .txt extension, the text specified is treated as the actual prompt.

### {{{{AI!`source-doc.docx`!`prompt_file.txt`!`section=section header&words=100`}}}}
Summarize a section of the document identified by 'section header' in the document located at '{self.ai_dir}/source-doc.docx'. The summary will be limited to 100 words or less. The prompt for the summary can be found in '{self.ai_dir}/prompt_file.txt', or the text provided directly if not a .txt file.

### {{{{AI!`source-doc.docx`!`prompt_file.txt`!`section=Attractions:Unique Experiences&words=100`}}}}
Summarize a range of content from 'Attractions' to 'Unique Experiences' in the document located at '{self.ai_dir}/source-doc.docx'. The summary will be limited to 100 words or less. The prompt for the summary can be found in '{self.ai_dir}/prompt_file.txt', or the text provided directly if not a .txt file.

An OpenAI API key is required for this feature and will be prompted when the application starts.
"""
        return help_text

    def _normalize_text(self, text):
        """Normalize text for section name comparison to handle apostrophe variations and special characters."""
        if not text:
            return ""
        
        # Convert to lowercase
        normalized = text.lower()
        
        # Replace various apostrophe types with a standard one (ASCII apostrophe)
        apostrophes = ["'", "'", "′", "‵", "´", "`", "′", "ʼ", "ʻ", "ʾ", "ʿ", "̓ ", "̕", "̔", "̛"]
        for apostrophe in apostrophes:
            if apostrophe != "'":  # Keep the standard apostrophe but replace others
                normalized = normalized.replace(apostrophe, "'")
        
        # Replace quotes with standard quotes
        quotes = [""", """, "„", "‟"]
        for quote in quotes:
            normalized = normalized.replace(quote, '"')
        
        # Remove other punctuation that might cause mismatches but preserve apostrophes
        for char in [',', '.', ':', ';', '!', '?', '-', '_', '(', ')', '[', ']', '{', '}', '/']:
            normalized = normalized.replace(char, " ")
        
        # Remove extra whitespace, including at start and end
        normalized = " ".join(normalized.split())
        
        return normalized 