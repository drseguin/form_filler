#!/usr/bin/env python3
# test_section.py - Test script for section extraction

import sys
import os
import logging
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from keyword_parser import keywordParser
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger('test_section')

def main():
    """Test section extraction from 1000-islands.docx"""
    logger.info("Starting section extraction test")
    
    # Path to test document
    doc_path = os.path.join('templates', '1000-islands.docx')
    if not os.path.exists(doc_path):
        logger.error(f"Test document not found: {doc_path}")
        return
    
    logger.info(f"Found test document: {doc_path}")
    
    # Create parser
    parser = keywordParser()
    
    # Analyze document structure
    doc = Document(doc_path)
    logger.info(f"Document contains {len(doc.paragraphs)} paragraphs")
    
    # List all paragraphs with styles
    logger.info("Document paragraph styles:")
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else "No style"
        text = para.text[:50] + "..." if len(para.text) > 50 else para.text
        if text.strip():  # Only log non-empty paragraphs
            logger.info(f"Paragraph {i}: Style='{style_name}', Text='{text}'")
    
    # Test extraction of 'Conclusion' section
    section_name = "Conclusion"
    logger.info(f"Testing extraction of section: '{section_name}'")
    
    # Modify the parser's state to simulate being in a Word document context
    parser.word_document = Document()
    
    # Process the template keyword to extract the section
    template_keyword = f"1000-islands.docx!section={section_name}"
    result = parser._process_template_keyword(template_keyword)
    
    if isinstance(result, dict) and "docx_template" in result:
        logger.info(f"Success! Section extracted to: {result['docx_template']}")
    else:
        logger.error(f"Section extraction failed. Result: {result}")

if __name__ == "__main__":
    main() 