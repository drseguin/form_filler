#!/usr/bin/env python
import sys
import os
import logging
import docx

# Add parent directory to path for imports
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from keyword_parser import keywordParser

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger("diagnose_template")

def diagnose_template_keyword(template_path, keyword_text):
    """
    Diagnose issues with template keywords, especially those with apostrophes or special characters.
    
    Args:
        template_path: Path to the Word document template
        keyword_text: The template keyword text to diagnose (e.g., "Millionaires' Row")
    """
    try:
        # Create a parser instance
        parser = keywordParser()
        
        # Log template file info
        if not os.path.exists(template_path):
            logger.error(f"Template file not found: {template_path}")
            return
        
        logger.info(f"Analyzing template file: {template_path}")
        doc = docx.Document(template_path)
        
        # Find all potential section headings
        headings = []
        for i, para in enumerate(doc.paragraphs):
            # Check if it looks like a heading or title
            is_heading = para.style and "heading" in para.style.name.lower()
            is_title = (para.text.strip() and 
                       len(para.text.strip()) < 100 and 
                       not para.text.strip().endswith('.') and
                       not para.text.strip().endswith(','))
            
            if is_heading or is_title:
                heading_text = para.text.strip()
                headings.append((i, heading_text))
                
                # Check if this might be the heading we're looking for
                if keyword_text.lower() in heading_text.lower() or heading_text.lower() in keyword_text.lower():
                    logger.info(f"\nPOTENTIAL MATCH at paragraph {i}: '{heading_text}'")
                    
                    # Show character details for debugging
                    logger.info(f"Characters in heading: {' '.join([f'{c}({ord(c)})' for c in heading_text])}")
                    
                    # Display exact differences using difflib
                    import difflib
                    diff = difflib.ndiff(keyword_text, heading_text)
                    logger.info(f"Diff between keyword and heading:")
                    logger.info('\n'.join(diff))
                    
                    # Try to extract with our normalized matching
                    normalized_keyword = parser._normalize_text(keyword_text)
                    normalized_heading = parser._normalize_text(heading_text)
                    
                    if normalized_keyword == normalized_heading:
                        logger.info("EXACT MATCH after normalization!")
                    elif normalized_keyword in normalized_heading:
                        logger.info("PARTIAL MATCH after normalization (keyword in heading)")
                    elif normalized_heading in normalized_keyword:
                        logger.info("PARTIAL MATCH after normalization (heading in keyword)")
                    else:
                        logger.info("NO MATCH after normalization")
                    
                    # Check for encoding issues
                    keyword_encoded = keyword_text.encode('utf-8')
                    heading_encoded = heading_text.encode('utf-8')
                    logger.info(f"Keyword bytes: {keyword_encoded}")
                    logger.info(f"Heading bytes: {heading_encoded}")
        
        # Test the extraction using the keyword
        logger.info("\nTesting section extraction:")
        filename = os.path.basename(template_path)
        template_keyword = f"{filename}!section={keyword_text}"
        
        logger.info(f"Using template keyword content: {template_keyword}")
        result = parser._process_template_keyword(template_keyword)
        
        if isinstance(result, dict) and "docx_template" in result:
            logger.info(f"SUCCESS: Section extracted to {result['docx_template']}")
            extracted_doc = docx.Document(result["docx_template"])
            logger.info(f"Extracted content has {len(extracted_doc.paragraphs)} paragraphs")
        else:
            logger.error(f"FAILED: {result}")
            
            # Suggest alternatives
            logger.info("\nSuggested alternatives to try:")
            for i, heading in headings:
                norm_heading = parser._normalize_text(heading)
                norm_keyword = parser._normalize_text(keyword_text)
                # Calculate similarity
                from difflib import SequenceMatcher
                similarity = SequenceMatcher(None, norm_heading, norm_keyword).ratio()
                if similarity > 0.5:  # Show reasonably similar headings
                    logger.info(f"  - Try: {heading} (similarity: {similarity:.2f})")
                    
            logger.info("\nOr try partial match with: {{TEMPLATE!%s!section=%s}}" % 
                       (os.path.basename(template_path), keyword_text.split()[0]))
    
    except Exception as e:
        logger.error(f"Error diagnosing template keyword: {str(e)}", exc_info=True)

def main():
    if len(sys.argv) < 3:
        print("Usage: python diagnose_template_keyword.py <template_file> <section_name>")
        print("Example: python diagnose_template_keyword.py templates/1000-islands.docx \"Millionaires' Row\"")
        return
    
    template_path = sys.argv[1]
    keyword_text = sys.argv[2]
    
    diagnose_template_keyword(template_path, keyword_text)

if __name__ == "__main__":
    main() 