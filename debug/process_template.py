"""Template processing implementation for keyword_parser.py"""

def _process_template_keyword(self, content):
    """Process template keywords using '!' separator."""
    if not content:
        return "[Invalid TEMPLATE reference]"

    try:
        # Split into filename and optional parameters using '!'
        parts = content.split("!")
        filename = parts[0].strip()
        
        self.logger.info(f"Processing TEMPLATE keyword with filename: '{filename}', parts: {parts}")

        # Handle library templates {{TEMPLATE!LIBRARY!template_name!version}}
        if filename.upper() == "LIBRARY":
            if len(parts) > 1:
                template_name = parts[1].strip() if len(parts) > 1 else ""
                template_version = parts[2].strip() if len(parts) > 2 else "DEFAULT"
                # Implement template library lookup here
                return f"[Template Library: {template_name} (Version: {template_version})]"
            return "[Invalid library template reference]"

        # Always look in the templates directory
        template_path = os.path.join('templates', filename)
        self.logger.info(f"Template path resolved to: {template_path}")

        # Check if file exists first
        if not os.path.exists(template_path):
            self.logger.warning(f"Template file not found: {template_path}")
            return f"[Template file not found: {template_path}]"
        
        self.logger.info(f"Template file exists: {template_path}")
        
        # Parse additional parameters
        param_part = parts[1] if len(parts) > 1 else ""
        has_section = param_part.startswith("section=")
        
        # For Word documents with section parameter
        if filename.lower().endswith('.docx') and has_section:
            # Parse section parameter
            section_info = self._parse_section_param(param_part)
            self.logger.info(f"Section info: {section_info}")
            
            start_section = section_info['start']
            end_section = section_info['end']
            
            if end_section:
                self.logger.info(f"Processing section range: '{start_section}' to '{end_section}'")
            else:
                self.logger.info(f"Processing single section: '{start_section}'")
                
            try:
                from docx import Document
                doc = Document(template_path)
                
                # Find sections and extract content
                found_start = False
                found_end = False
                section_start_index = -1
                section_end_index = -1
                
                # Map of all headings for debugging
                headings = []
                
                # First pass: look for start and end by headings
                for i, para in enumerate(doc.paragraphs):
                    # Check if this is a heading-like paragraph
                    is_heading = para.style and "heading" in para.style.name.lower()
                    is_title = (para.text.strip() and 
                                len(para.text.strip()) < 100 and 
                                not para.text.strip().endswith('.') and
                                not para.text.strip().endswith(','))
                    
                    if is_heading or is_title:
                        headings.append((i, para.text.strip()))
                        
                        # Look for start section
                        if not found_start and (start_section.lower() == para.text.strip().lower() or
                                              start_section.lower() in para.text.strip().lower()):
                            found_start = True
                            section_start_index = i + 1  # Start after this heading
                            self.logger.info(f"Found start section at paragraph {i}: '{para.text}'")
                        
                        # Look for end section if specified
                        elif found_start and end_section and (end_section.lower() == para.text.strip().lower() or
                                                           end_section.lower() in para.text.strip().lower()):
                            found_end = True
                            section_end_index = i  # End before this heading
                            self.logger.info(f"Found end section at paragraph {i}: '{para.text}'")
                            break
                        
                        # If we've found the start and no specific end is required,
                        # any next heading marks the end of the section
                        elif found_start and not end_section:
                            found_end = True
                            section_end_index = i
                            self.logger.info(f"Found next heading at paragraph {i}: '{para.text}'")
                            break
                
                # If we didn't find start by heading match, try exact text match
                if not found_start:
                    self.logger.info(f"Looking for exact text match for start section: '{start_section}'")
                    for i, para in enumerate(doc.paragraphs):
                        if start_section.lower() == para.text.strip().lower():
                            found_start = True
                            section_start_index = i + 1  # Start after this paragraph
                            self.logger.info(f"Found start section by exact match at paragraph {i}")
                            break
                
                # If we found start but not end, and end is specified, look for exact match
                if found_start and not found_end and end_section:
                    self.logger.info(f"Looking for exact text match for end section: '{end_section}'")
                    for i in range(section_start_index, len(doc.paragraphs)):
                        if end_section.lower() == doc.paragraphs[i].text.strip().lower():
                            found_end = True
                            section_end_index = i
                            self.logger.info(f"Found end section by exact match at paragraph {i}")
                            break
                
                # If we found start but not end, use end of document
                if found_start and not found_end:
                    section_end_index = len(doc.paragraphs)
                    self.logger.info(f"Using end of document as section end (paragraph {section_end_index})")
                
                # Log error if section not found
                if not found_start:
                    self.logger.warning(f"Could not find section '{start_section}'")
                    if headings:
                        self.logger.info("Available headings:")
                        for idx, heading in headings:
                            self.logger.info(f"  Paragraph {idx}: '{heading}'")
                    return f"[Section '{start_section}' not found in {filename}]"
                
                # Extract the selected paragraphs
                section_paragraphs = doc.paragraphs[section_start_index:section_end_index]
                if not section_paragraphs:
                    self.logger.warning(f"No content found in section")
                    return f"[No content found in section]"
                    
                self.logger.info(f"Extracted {len(section_paragraphs)} paragraphs")
                
                # Create a new document with just the section content
                if self.word_document:
                    temp_doc = Document()
                    
                    # Add title with section name
                    title_para = temp_doc.add_paragraph(start_section)
                    try:
                        title_para.style = 'Heading 1'
                    except:
                        # Manually style if needed
                        title_run = title_para.runs[0]
                        title_run.bold = True
                        title_run.font.size = Pt(16)
                    
                    # Copy all paragraphs with formatting
                    for para in section_paragraphs:
                        p = temp_doc.add_paragraph()
                        # Copy text and formatting
                        for run in para.runs:
                            r = p.add_run(run.text)
                            r.bold = run.bold
                            r.italic = run.italic
                            r.underline = run.underline
                            if run.font.size:
                                r.font.size = run.font.size
                            if run.font.name:
                                r.font.name = run.font.name
                            if run.font.color.rgb:
                                r.font.color.rgb = run.font.color.rgb
                        
                        # Copy paragraph formatting
                        try:
                            if para.style:
                                p.style = para.style.name
                            p.paragraph_format.alignment = para.paragraph_format.alignment
                            p.paragraph_format.left_indent = para.paragraph_format.left_indent
                            p.paragraph_format.right_indent = para.paragraph_format.right_indent
                            p.paragraph_format.space_before = para.paragraph_format.space_before
                            p.paragraph_format.space_after = para.paragraph_format.space_after
                        except:
                            pass
                    
                    # Save to temporary file
                    import tempfile
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                        section_path = tmp.name
                        temp_doc.save(section_path)
                    
                    self.logger.info(f"Created section document at {section_path}")
                    return {"docx_template": section_path}
                else:
                    # Return as plain text if not in a Word context
                    return "\n\n".join([para.text for para in section_paragraphs])
                
            except ImportError:
                self.logger.error("python-docx library not available")
                return f"[Error: python-docx library not available]"
            except Exception as e:
                self.logger.error(f"Error extracting section: {str(e)}", exc_info=True)
                return f"[Error extracting section: {str(e)}]"
        
        # For Word documents (whole document) 
        elif filename.lower().endswith('.docx') and self.word_document:
            self.logger.info(f"Returning whole Word document template")
            return {"docx_template": template_path}
            
        # For non-Word files or other parameters
        else:
            try:
                # Read file content
                if filename.lower().endswith('.docx'):
                    try:
                        from docx import Document
                        doc = Document(template_path)
                        file_content = "\n\n".join([para.text for para in doc.paragraphs])
                        # Extract tables
                        for table in doc.tables:
                            table_text = []
                            for row in table.rows:
                                row_text = []
                                for cell in row.cells:
                                    row_text.append(cell.text)
                                table_text.append(" | ".join(row_text))
                            file_content += "\n\n" + "\n".join(table_text)
                    except ImportError:
                        return "[Error: python-docx library not available]"
                else:
                    with open(template_path, 'r', encoding='utf-8') as file:
                        file_content = file.read()
                
                # Process special parameters
                if param_part:
                    if param_part.startswith("line="):
                        # Get specific line
                        try:
                            line_num = int(param_part.split("line=")[1].split(",")[0].strip())
                            lines = file_content.splitlines()
                            if 0 <= line_num - 1 < len(lines):
                                return lines[line_num - 1]
                            return f"[Line {line_num} not found]"
                        except ValueError:
                            return f"[Invalid line number]"
                    
                    elif param_part.startswith("paragraph="):
                        # Get specific paragraph
                        try:
                            para_num = int(param_part.split("paragraph=")[1].split(",")[0].strip())
                            paras = file_content.split("\n\n")
                            if 0 <= para_num - 1 < len(paras):
                                return paras[para_num - 1]
                            return f"[Paragraph {para_num} not found]"
                        except ValueError:
                            return f"[Invalid paragraph number]"
                    
                    elif param_part.startswith("section=") and not filename.lower().endswith('.docx'):
                        # Text file section extraction
                        section_info = self._parse_section_param(param_part)
                        section_name = section_info['start']
                        lines = file_content.splitlines()
                        section_lines = []
                        found_section = False
                        
                        for line in lines:
                            if not found_section:
                                if line.strip().lower() == section_name.lower() or line.strip().lower() == f"{section_name.lower()}:":
                                    found_section = True
                                    continue
                            else:
                                if line.strip() and not line.startswith(' ') and line.strip().endswith(':'):
                                    # Next section found
                                    break
                                section_lines.append(line)
                        
                        if found_section:
                            return "\n".join(section_lines)
                        else:
                            return f"[Section {section_name} not found]"
                    
                    elif param_part.startswith("VARS("):
                        # Template with variable substitution
                        try:
                            vars_text = param_part.split("VARS(")[1].split(")")[0]
                            var_pairs = vars_text.split(",")
                            
                            # Parse variables
                            variables = {}
                            for pair in var_pairs:
                                if "=" in pair:
                                    key, value = pair.split("=", 1)
                                    # Parse values that might be keywords
                                    variables[key.strip()] = self.parse(value.strip())
                            
                            # Replace in template
                            result = file_content
                            for key, value in variables.items():
                                result = result.replace(f"{{{key}}}", str(value) if value is not None else "")
                            
                            return result
                        except Exception as e:
                            return f"[Error in template variables: {str(e)}]"
                
                # Return the whole file if no special parameters
                return file_content
                
            except Exception as e:
                self.logger.error(f"Error processing template: {str(e)}", exc_info=True)
                return f"[Error: {str(e)}]"
    
    except Exception as e:
        self.logger.error(f"Error processing template: {str(e)}", exc_info=True)
        return f"[Template error: {str(e)}]" 