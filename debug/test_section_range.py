#!/usr/bin/env python3
# test_section_range.py - Test script for section range extraction

import sys
import os
import logging
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from keyword_parser import keywordParser
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger('test_section_range')

def main():
    """Test section range extraction from 1000-islands.docx"""
    logger.info("Starting section range extraction test")
    
    # Path to test document
    doc_path = os.path.join('templates', '1000-islands.docx')
    if not os.path.exists(doc_path):
        logger.error(f"Test document not found: {doc_path}")
        return
    
    logger.info(f"Found test document: {doc_path}")
    
    # Create parser
    parser = keywordParser()
    
    # Analyze document structure to show available sections
    doc = Document(doc_path)
    logger.info(f"Document contains {len(doc.paragraphs)} paragraphs")
    
    # List potential section headings
    logger.info("Potential section headings:")
    for i, para in enumerate(doc.paragraphs):
        is_heading = para.style and "heading" in para.style.name.lower()
        is_standalone = (para.text.strip() and 
                         len(para.text.strip()) < 50 and 
                         not para.text.strip().endswith('.') and
                         not para.text.strip().endswith(','))
        
        if is_heading or is_standalone:
            logger.info(f"Paragraph {i}: '{para.text}'")
    
    # Test extraction of section range
    start_section = "Lost Villages"
    end_section = "Millionaires' Row"
    logger.info(f"Testing extraction of section range: '{start_section}' to '{end_section}'")
    
    # Modify the parser's state to simulate being in a Word document context
    parser.word_document = Document()
    
    # Process the template keyword with section range
    template_keyword = f"1000-islands.docx!section={start_section}:{end_section}"
    result = parser._process_template_keyword(template_keyword)
    
    # Check the result
    if isinstance(result, dict) and "docx_template" in result:
        logger.info(f"Success! Section range extracted to: {result['docx_template']}")
        
        # Inspect the extracted content
        extracted_doc = Document(result['docx_template'])
        logger.info(f"Extracted document contains {len(extracted_doc.paragraphs)} paragraphs")
        
        # Show first few paragraphs of extracted content
        for i, para in enumerate(extracted_doc.paragraphs[:5]):
            if para.text.strip():
                text = para.text[:50] + "..." if len(para.text) > 50 else para.text
                logger.info(f"Extracted paragraph {i}: '{text}'")
    else:
        logger.error(f"Section range extraction failed. Result: {result}")

if __name__ == "__main__":
    main() 