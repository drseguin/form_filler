#!/usr/bin/env python
import sys
import os
import logging
import docx

# Add parent directory to path for imports
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from keyword_parser import keywordParser

# Set up logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
                   handlers=[logging.StreamHandler()])

logger = logging.getLogger("template_test")

def list_document_sections(doc_path):
    """List all potential sections in a Word document"""
    doc = docx.Document(doc_path)
    
    logger.info(f"Examining document: {doc_path}")
    logger.info("Potential section headings:")
    
    for i, para in enumerate(doc.paragraphs):
        # Check if this looks like a heading or title
        is_heading = para.style and "heading" in para.style.name.lower()
        is_title = (para.text.strip() and 
                   len(para.text.strip()) < 100 and 
                   not para.text.strip().endswith('.') and
                   not para.text.strip().endswith(','))
        
        if is_heading or is_title:
            style_name = para.style.name if para.style else "No style"
            logger.info(f"  [{i}] '{para.text.strip()}' (Style: {style_name})")
            
            # Special case for Millionaires' Row - show more details
            if "millionaire" in para.text.lower():
                logger.info(f"  === FOUND MILLIONAIRES' ROW section ===")
                logger.info(f"  Exact text: '{para.text}'")
                for run in para.runs:
                    logger.info(f"  Run text: '{run.text}', Bold: {run.bold}, Italic: {run.italic}")
                    
    return doc.paragraphs

def test_template_extraction(template_path, section_name):
    """Test template section extraction using keywordParser"""
    parser = keywordParser()
    
    # Get just the filename from the path
    filename = os.path.basename(template_path)
    
    # Correctly construct the keyword content (without the outer braces)
    template_keyword_content = f"{filename}!section={section_name}"
    logger.info(f"Testing template keyword: {{{{{template_keyword_content}}}}}")
    
    # Process the template keyword - note we pass only the content (not the TEMPLATE! prefix)
    result = parser._process_template_keyword(template_keyword_content)
    
    # Check the result
    if isinstance(result, dict) and "docx_template" in result:
        logger.info(f"SUCCESS: Section extracted to {result['docx_template']}")
        # Verify the content
        extracted_doc = docx.Document(result["docx_template"])
        logger.info(f"Extracted content has {len(extracted_doc.paragraphs)} paragraphs")
        for i, para in enumerate(extracted_doc.paragraphs):
            if i < 3 or i >= len(extracted_doc.paragraphs) - 3:  # Show first and last few paragraphs
                logger.info(f"  [{i}] {para.text[:100]}...")
    else:
        logger.error(f"FAILED: {result}")

def main():
    template_path = "templates/1000-islands.docx"
    
    # First, list all sections
    paragraphs = list_document_sections(template_path)
    
    # Test extraction with different apostrophe variants
    logger.info("\nTesting with different apostrophe formats for Millionaires' Row:")
    test_template_extraction(template_path, "Millionaires' Row")  # Standard ASCII apostrophe
    test_template_extraction(template_path, "Millionaires' Row")  # Unicode right single quotation mark
    test_template_extraction(template_path, "Millionaires Row")   # Without apostrophe
    
    # Try using partial matching
    logger.info("\nTesting with partial section name:")
    test_template_extraction(template_path, "Millionaires")
    
    # Also test a section we know exists
    logger.info("\nTesting with a known section:")
    test_template_extraction(template_path, "Introduction")

if __name__ == "__main__":
    main() 