# main.py
import streamlit as st
import os
import re
import docx
import tempfile
import time
import logging
from excel_manager import excelManager # Assuming excel_manager.py is in the same directory
from keyword_parser import keywordParser # Assuming keyword_parser.py is in the same directory
from collections import Counter
from logs.logger_config import setup_logger
from pathlib import Path

# Setup logger
logger = setup_logger('main')

def check_openai_api_key() -> bool:
    """
    Check if the OpenAI API key is set in the session state or in the .streamlit/secrets.toml file.
    
    Returns:
        bool: True if the API key is set, False otherwise
    """
    # First check if API key is in session state
    if 'openai_api_key' in st.session_state and st.session_state['openai_api_key']:
        logger.info("OpenAI API key found in session state")
        return True
        
    # If not in session state, check secrets.toml
    secrets_path = Path(".streamlit/secrets.toml")
    
    if not secrets_path.exists():
        logger.debug("Secrets file not found: .streamlit/secrets.toml")
        return False
        
    try:
        # Parse toml file manually
        with open(secrets_path, 'r', encoding='utf-8') as file:
            for line in file:
                if line.strip().startswith('openai_api_key'):
                    parts = line.strip().split('=', 1)
                    if len(parts) == 2:
                        api_key = parts[1].strip().strip('"\'')
                        if api_key:
                            logger.info("OpenAI API key found in secrets.toml")
                            # Store in session state for future use
                            st.session_state['openai_api_key'] = api_key
                            return True
                        else:
                            logger.warning("OpenAI API key is empty in .streamlit/secrets.toml")
                            return False
    except Exception as e:
        logger.error(f"Error reading secrets file: {str(e)}", exc_info=True)
        return False
    
    logger.warning("OpenAI API key not found in session state or secrets.toml")
    return False

# Function to get the API key (for use in OpenAI client)
def get_openai_api_key() -> str:
    """
    Get the OpenAI API key from session state or return empty string if not set.
    
    Returns:
        str: The OpenAI API key or empty string
    """
    return st.session_state.get('openai_api_key', '')

def preprocess_word_doc(doc_path):
    """
    Analyze a Word document to determine what keywords it contains, using '!' separator.

    Args:
        doc_path: Path to the Word document

    Returns:
        Dictionary with keyword counts and whether Excel file is needed
    """
    logger.info(f"Preprocessing Word document: {doc_path}")
    doc = docx.Document(doc_path)
    pattern = r'{{(.*?)}}'

    keywords = {
        "excel": {"CELL": [], "LAST": [], "RANGE": [], "COLUMN": [], "OTHER": []},
        "input": {"text": [], "area": [], "date": [], "select": [], "check": []},
        "template": {"full": [], "section": [], "range": []},
        "json": [],
        "ai": [],
        "other": []
    }
    needs_excel = False
    needs_templates = False
    needs_json = False
    needs_ai = False
    total_keywords = 0
    excel_files = set()  # Store unique Excel files needed
    excel_files_not_found = []  # Store Excel files that were not found
    template_files = set()  # Store unique template files needed
    template_files_not_found = []  # Store template files that were not found
    json_files = set()  # Store unique JSON files needed
    json_files_not_found = []  # Store JSON files that were not found
    ai_source_files = set()  # Store unique AI source files needed
    ai_source_files_not_found = []  # Store AI source files that were not found
    ai_prompt_files = set()  # Store unique AI prompt files needed
    ai_prompt_files_not_found = []  # Store AI prompt files that were not found

    # Ensure excel directory exists
    excel_dir = "excel"
    if not os.path.exists(excel_dir):
        os.makedirs(excel_dir)
        logger.info(f"Created excel directory: {excel_dir}")
        
    # Ensure templates directory exists
    templates_dir = "templates"
    if not os.path.exists(templates_dir):
        os.makedirs(templates_dir)
        logger.info(f"Created templates directory: {templates_dir}")
        
    # Ensure json directory exists
    json_dir = "json"
    if not os.path.exists(json_dir):
        os.makedirs(json_dir)
        logger.info(f"Created json directory: {json_dir}")
        
    # Ensure ai directory exists
    ai_dir = "ai"
    if not os.path.exists(ai_dir):
        os.makedirs(ai_dir)
        logger.info(f"Created ai directory: {ai_dir}")

    def categorize_keyword(content):
        nonlocal needs_excel
        nonlocal needs_templates
        nonlocal needs_json
        nonlocal needs_ai
        parts = content.split("!", 1) # Use '!' separator
        keyword_type = parts[0].strip().upper()

        if not keyword_type: return # Ignore empty keywords {{}}

        if keyword_type == "XL":
            needs_excel = True
            
            if len(parts) > 1:
                # Check if the next part might be an Excel file path
                excel_parts = parts[1].split("!", 1)
                excel_file = excel_parts[0].strip()
                rest_of_content = excel_parts[1] if len(excel_parts) > 1 else ""
                
                # If the excel_file looks like a file path ending with .xlsx or .xls
                if excel_file.lower().endswith(('.xlsx', '.xls')):
                    excel_files.add(excel_file)  # Add to set of required Excel files
                    
                    # Check if file exists in current path or excel folder
                    file_exists = os.path.exists(excel_file) or os.path.exists(os.path.join(excel_dir, excel_file))
                    
                    if not file_exists and excel_file not in excel_files_not_found:
                        excel_files_not_found.append(excel_file)
                        logger.info(f"Excel file not found: {excel_file}")
                    
                    # Re-parse the rest as a normal Excel keyword
                    if rest_of_content:
                        # Check if the next part is a valid Excel subtype
                        sub_parts = rest_of_content.split("!", 1)
                        xl_subtype = sub_parts[0].strip().upper()
                        if xl_subtype in keywords["excel"]:
                            keywords["excel"][xl_subtype].append(content)
                        else:
                            keywords["excel"]["OTHER"].append(content)
                    else:
                        keywords["excel"]["OTHER"].append(content)
                else:
                    # Old format or just XL!SUBTYPE without an Excel file specified
                    sub_parts = parts[1].split("!", 1)
                    xl_subtype = sub_parts[0].strip().upper()
                    if xl_subtype in keywords["excel"]:
                        keywords["excel"][xl_subtype].append(content)
                    else:
                        # If subtype unknown, check if it looks like an old format/named range
                        if ':' not in parts[1] and '!' not in parts[1]: # Likely named range or old cell ref
                            keywords["excel"]["RANGE"].append(content) # Assume RANGE for named range
                        else:
                            keywords["excel"]["OTHER"].append(content) # Potentially old or invalid format
            else:
                keywords["excel"]["OTHER"].append(content) # Invalid XL format {{XL}}

        elif keyword_type == "INPUT":
            if len(parts) > 1:
                input_parts = parts[1].split("!")
                input_type = input_parts[0].lower() if input_parts else "text"
                if input_type in keywords["input"]:
                    keywords["input"][input_type].append(content)
                else:
                    keywords["input"]["text"].append(content)
            else:
                 keywords["input"]["text"].append(content) # {{INPUT}} defaults to text

        elif keyword_type == "TEMPLATE":
            needs_templates = True
            # Extract sections/ranges correctly
            if len(parts) <= 1:
                # No parameters, just the keyword type - should never happen for TEMPLATE but handle it
                logger.info(f"Categorizing bare TEMPLATE keyword without parameters")
                keywords["template"]["full"].append(content)
                return
                
            # Split the content after the TEMPLATE! prefix to analyze the parts
            template_parts = parts[1].split("!")
            template_path = template_parts[0]
            
            logger.info(f"Processing TEMPLATE keyword: '{content}' with path '{template_path}'")
            
            # Add template file to the set of required templates
            if template_path.lower().endswith(('.docx', '.txt')):
                template_files.add(template_path)
                
                # Check if file exists in current path or templates folder
                file_exists = os.path.exists(template_path) or os.path.exists(os.path.join(templates_dir, template_path))
                
                if not file_exists and template_path not in template_files_not_found:
                    template_files_not_found.append(template_path)
                    logger.info(f"Template file not found: {template_path}")
            
            # If there's no second part with section=, it's a full template
            if len(template_parts) == 1:
                logger.info(f"Categorizing as FULL template: {content}")
                keywords["template"]["full"].append(content)
            # Check for section parameter
            elif len(template_parts) > 1 and "section=" in template_parts[1]:
                # Need to extract the section value to check for colon
                try:
                    section_param = template_parts[1].split("section=")[1].split("&")[0] if "section=" in template_parts[1] else ""
                    logger.info(f"Found section parameter: '{section_param}'")
                    
                    # Check if it's a range (contains ':') - {{TEMPLATE!filename.docx!section=Start:End}}
                    if ":" in section_param:
                        logger.info(f"Categorizing as RANGE template: {content}")
                        keywords["template"]["range"].append(content)
                    else:
                        # Just a single section - {{TEMPLATE!filename.docx!section=SectionName}}
                        logger.info(f"Categorizing as SECTION template: {content}")
                        keywords["template"]["section"].append(content)
                except Exception as e:
                    logger.error(f"Error parsing section parameter: {e}")
                    # Default to full if we can't parse the section
                    keywords["template"]["full"].append(content)
            else:
                # Any other template format defaults to full template
                logger.info(f"Categorizing as FULL template (default): {content}")
                keywords["template"]["full"].append(content)
        elif keyword_type == "JSON":
            needs_json = True
            keywords["json"].append(content)
            
            if len(parts) > 1:
                # Extract the JSON file name
                json_parts = parts[1].split("!", 1)
                json_file = json_parts[0].strip()
                
                # Handle special case where the first part might be empty ({{JSON!!filename.json}})
                if not json_file and len(json_parts) > 1:
                    # In this case, the second part is the filename
                    json_file = json_parts[1].split("!", 1)[0].strip()
                
                # If the json_file looks like a file path ending with .json
                if json_file.lower().endswith('.json'):
                    json_files.add(json_file)  # Add to set of required JSON files
                    
                    # Check if file exists in current path or json folder
                    file_exists = os.path.exists(json_file) or os.path.exists(os.path.join(json_dir, json_file))
                    
                    if not file_exists and json_file not in json_files_not_found:
                        json_files_not_found.append(json_file)
                        logger.info(f"JSON file not found: {json_file}")
            
        elif keyword_type == "AI":
            needs_ai = True
            keywords["ai"].append(content)
            
            if len(parts) > 1:
                # Extract the AI source document and prompt files
                ai_parts = parts[1].split("!")
                
                # First part is always the source document
                if len(ai_parts) >= 1:
                    source_file = ai_parts[0].strip()
                    if source_file.lower().endswith(('.docx', '.txt')):
                        ai_source_files.add(source_file)
                        
                        # Check if file exists in current path or ai folder
                        file_exists = os.path.exists(source_file) or os.path.exists(os.path.join(ai_dir, source_file))
                        
                        if not file_exists and source_file not in ai_source_files_not_found:
                            ai_source_files_not_found.append(source_file)
                            logger.info(f"AI source file not found: {source_file}")
                
                # Second part could be a prompt file or a literal prompt
                if len(ai_parts) >= 2:
                    prompt = ai_parts[1].strip()
                    # Only treat as a file if it ends with .txt
                    if prompt.lower().endswith('.txt'):
                        ai_prompt_files.add(prompt)
                        
                        # Check if file exists in current path or ai folder
                        file_exists = os.path.exists(prompt) or os.path.exists(os.path.join(ai_dir, prompt))
                        
                        if not file_exists and prompt not in ai_prompt_files_not_found:
                            ai_prompt_files_not_found.append(prompt)
                            logger.info(f"AI prompt file not found: {prompt}")
        else:
             # If not a recognized type, check if it might be an Excel named range
             if '!' not in content and ':' not in content:
                  needs_excel = True
                  keywords["excel"]["RANGE"].append(content) # Treat as potential named range
             else:
                  keywords["other"].append(content)

    # Scan paragraphs
    for paragraph in doc.paragraphs:
        matches = list(re.finditer(pattern, paragraph.text))
        total_keywords += len(matches)
        for match in matches:
            categorize_keyword(match.group(1))

    # Scan tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches = list(re.finditer(pattern, paragraph.text))
                    total_keywords += len(matches)
                    for match in matches:
                        categorize_keyword(match.group(1))

    summary = {
        "total_keywords": total_keywords,
        "excel_counts": {k: len(v) for k, v in keywords["excel"].items()},
        "input_counts": {k: len(v) for k, v in keywords["input"].items()},
        "template_count": {k: len(v) for k, v in keywords["template"].items()},
        "template_total": sum(len(v) for v in keywords["template"].values()),
        "json_count": len(keywords["json"]),
        "ai_count": len(keywords["ai"]),
        "other_count": len(keywords["other"]),
        "needs_excel": needs_excel,
        "needs_templates": needs_templates,
        "needs_json": needs_json,
        "needs_ai": needs_ai,
        "excel_files": list(excel_files),
        "excel_files_not_found": excel_files_not_found,
        "template_files": list(template_files),
        "template_files_not_found": template_files_not_found,
        "json_files": list(json_files),
        "json_files_not_found": json_files_not_found,
        "ai_source_files": list(ai_source_files),
        "ai_source_files_not_found": ai_source_files_not_found,
        "ai_prompt_files": list(ai_prompt_files),
        "ai_prompt_files_not_found": ai_prompt_files_not_found,
        "keywords": keywords
    }
    
    # Debug log for template counts
    logger.info(f"Template summary: {summary['template_count']}")
    logger.info(f"Template total: {summary['template_total']}")
    for t_type, items in keywords["template"].items():
        if items:
            logger.info(f"Template {t_type} items: {items}")
    
    # Debug log for Excel files
    if excel_files:
        logger.info(f"Excel files needed: {list(excel_files)}")
    if excel_files_not_found:
        logger.info(f"Excel files not found: {excel_files_not_found}")
        
    # Debug log for Template files
    if template_files:
        logger.info(f"Template files needed: {list(template_files)}")
    if template_files_not_found:
        logger.info(f"Template files not found: {template_files_not_found}")
        
    # Debug log for JSON files
    if json_files:
        logger.info(f"JSON files needed: {list(json_files)}")
    if json_files_not_found:
        logger.info(f"JSON files not found: {json_files_not_found}")
        
    # Debug log for AI files
    if ai_source_files:
        logger.info(f"AI source files needed: {list(ai_source_files)}")
    if ai_source_files_not_found:
        logger.info(f"AI source files not found: {ai_source_files_not_found}")
    if ai_prompt_files:
        logger.info(f"AI prompt files needed: {list(ai_prompt_files)}")
    if ai_prompt_files_not_found:
        logger.info(f"AI prompt files not found: {ai_prompt_files_not_found}")
    
    return summary


def process_word_doc(doc_path, excel_path=None, parser=None):
    """
    Process a Word document, replacing keywords with values using the provided parser.

    Args:
        doc_path: Path to the Word document
        excel_path: Path to the Excel spreadsheet (optional - manager passed via parser)
        parser: An initialized keywordParser instance

    Returns:
        Processed document object and a count of replaced keywords
    """
    logger.info(f"Starting document processing: {doc_path}")
    if excel_path:
        logger.info(f"Using Excel file: {excel_path}")
        
    if not parser:
        raise ValueError("KeywordParser instance is required.")

    doc = docx.Document(doc_path)
    parser.set_word_document(doc) # Ensure parser has the correct document object

    pattern = r'{{(.*?)}}'
    total_keywords_initial = 0

    # Count initial keywords
    elements_to_scan = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                elements_to_scan.extend(cell.paragraphs)

    for paragraph in elements_to_scan:
        total_keywords_initial += len(re.findall(pattern, paragraph.text))

    logger.info(f"Found {total_keywords_initial} keywords in document")
    
    if total_keywords_initial == 0:
        st.warning("No keywords found in the document.")
        logger.warning("No keywords found in the document")
        return doc, 0

    progress_bar = st.progress(0)
    progress_text = st.empty()
    progress_text.text("Processing keywords...")

    # Process paragraph by paragraph, letting the parser handle replacements
    processed_keywords_count = 0
    elements_processed = 0
    total_elements = len(elements_to_scan)

    # Important: We'll set form_submitted to True BUT also ensure the input values
    # from the main form are properly transferred to the parser
    parser.form_submitted = True
    
    # Make sure all the input values from our main form are in the parser's input_values dict
    if 'input_values_main' in st.session_state and st.session_state.input_values_main:
        for content, value in st.session_state.input_values_main.items():
            # Store with the exact content format 
            keyword = f"{{{{{content}}}}}"
            parser.input_values[keyword] = value
            
            # Also store in alternate formats
            if content.startswith("INPUT!"):
                # Store without the INPUT! prefix
                non_prefix_content = content[6:]  # Remove "INPUT!"
                alt_keyword = f"{{{{{non_prefix_content}}}}}"
                parser.input_values[alt_keyword] = value
            else:
                # Store with the INPUT! prefix
                alt_keyword = f"{{{{INPUT!{content}}}}}"
                parser.input_values[alt_keyword] = value
    
    # Also check for any keywords directly in the form fields format
    form_keys = [k for k in st.session_state.keys() if k.startswith('input_field_')]
    if form_keys:
        for form_key in form_keys:
            if form_key.startswith('input_field_INPUT!'):
                content = form_key[12:]  # Remove 'input_field_' prefix
                value = st.session_state[form_key]
                keyword = f"{{{{{content}}}}}"
                parser.input_values[keyword] = value
    
    # Store input values for potential troubleshooting
    st.session_state['debug_input_values'] = parser.input_values.copy()

    for paragraph in elements_to_scan:
        original_text = paragraph.text
        keywords_in_para = len(re.findall(pattern, original_text))

        if keywords_in_para > 0:
            # Extract keywords in this paragraph for display
            keywords_in_this_para = re.findall(pattern, original_text)
            current_keyword = keywords_in_this_para[0] if keywords_in_this_para else "Unknown"
            
            # Update progress text to show current keyword
            progress_text.text(f"{processed_keywords_count}/{total_keywords_initial} - {{{{{current_keyword}}}}}")
            
            try:
                # parser.parse will handle replacements, including potential table creation
                parsed_result = parser.parse(original_text)

                # Check if we got a dict with a docx template
                if isinstance(parsed_result, dict) and "docx_template" in parsed_result:
                    try:
                        # Update the paragraph's text with any text content
                        paragraph.text = parsed_result["text"]
                        
                        # Get the template path from our result
                        template_path = parsed_result["docx_template"]
                        
                        # Load the template document with proper formatting
                        template_doc = docx.Document(template_path)
                        
                        # Insert the template document at the current paragraph location
                        paragraph_element = paragraph._element
                        paragraph_parent = paragraph_element.getparent()
                        paragraph_index = paragraph_parent.index(paragraph_element)
                        
                        # For each paragraph in the template, add it to the main document
                        for p in template_doc.paragraphs:
                            # Create a new paragraph in the main document
                            new_p = doc.add_paragraph()
                            # Copy over the paragraph's runs with their formatting
                            for run in p.runs:
                                new_run = new_p.add_run(run.text)
                                # Copy formatting from the original run
                                new_run.bold = run.bold
                                new_run.italic = run.italic
                                new_run.underline = run.underline
                                if run.font.size:
                                    new_run.font.size = run.font.size
                                if run.font.name:
                                    new_run.font.name = run.font.name
                                if run.font.color.rgb:
                                    new_run.font.color.rgb = run.font.color.rgb
                            
                            # Copy paragraph formatting
                            if p.style:
                                try:
                                    new_p.style = p.style.name
                                except:
                                    pass  # Style might not exist in target document
                            new_p.paragraph_format.alignment = p.paragraph_format.alignment
                            new_p.paragraph_format.left_indent = p.paragraph_format.left_indent
                            new_p.paragraph_format.right_indent = p.paragraph_format.right_indent
                            new_p.paragraph_format.space_before = p.paragraph_format.space_before
                            new_p.paragraph_format.space_after = p.paragraph_format.space_after
                            
                            # Position the new paragraph in the right place
                            new_p_element = new_p._element
                            # Remove it from wherever it was added automatically
                            new_p_element.getparent().remove(new_p_element)
                            # Insert at the right position
                            paragraph_index += 1
                            paragraph_parent.insert(paragraph_index, new_p_element)
                        
                        # For each table in the template, add it to the main document
                        for table in template_doc.tables:
                            # Create a new table in the main document with the same dimensions
                            rows = len(table.rows)
                            cols = len(table.columns)
                            new_table = doc.add_table(rows=rows, cols=cols)
                            
                            # Copy table style if it exists
                            if table.style:
                                try:
                                    new_table.style = table.style
                                except:
                                    pass  # Style might not exist in target document
                            
                            # Copy cell by cell
                            for i, row in enumerate(table.rows):
                                for j, cell in enumerate(row.cells):
                                    target_cell = new_table.cell(i, j)
                                    # Clear any default paragraph
                                    for p in target_cell.paragraphs:
                                        p._element.getparent().remove(p._element)
                                    
                                    # Copy each paragraph from source cell
                                    for p in cell.paragraphs:
                                        cell_p = target_cell.add_paragraph()
                                        for run in p.runs:
                                            cell_run = cell_p.add_run(run.text)
                                            cell_run.bold = run.bold
                                            cell_run.italic = run.italic
                                            cell_run.underline = run.underline
                                            if run.font.size:
                                                cell_run.font.size = run.font.size
                                            if run.font.name:
                                                cell_run.font.name = run.font.name
                                            if run.font.color.rgb:
                                                cell_run.font.color.rgb = run.font.color.rgb
                                            
                                        if p.style:
                                            try:
                                                cell_p.style = p.style.name
                                            except:
                                                pass
                                        cell_p.paragraph_format.alignment = p.paragraph_format.alignment
                            
                            # Position the new table in the right place
                            new_table_element = new_table._element
                            # Remove it from wherever it was added automatically
                            new_table_element.getparent().remove(new_table_element)
                            # Insert at the right position
                            paragraph_index += 1
                            paragraph_parent.insert(paragraph_index, new_table_element)
                            
                            # Add a blank paragraph after the table to prevent merging
                            blank_p = doc.add_paragraph()
                            blank_p_element = blank_p._element
                            blank_p_element.getparent().remove(blank_p_element)
                            paragraph_index += 1
                            paragraph_parent.insert(paragraph_index, blank_p_element)
                        
                        # Count as processed (we processed the entire keyword that produced the template)
                        processed_keywords_count += 1
                    except Exception as e:
                        st.error(f"Error inserting template document: {str(e)}")
                        logger.error(f"Error inserting template document: {str(e)}", exc_info=True)
                        # If template insertion fails, keep original text
                        paragraph.text = original_text

                # Check if we got a dict with a table object
                elif isinstance(parsed_result, dict) and "table" in parsed_result:
                    try:
                        # Update the paragraph's text with any text content
                        paragraph.text = parsed_result["text"]
                        
                        # Get the table object from our result
                        table = parsed_result["table"]
                        
                        # Insert the table at the current paragraph location
                        # We need to get the paragraph parent and insert after this paragraph
                        paragraph_element = paragraph._element
                        paragraph_parent = paragraph_element.getparent()
                        
                        # Insert the table after the current paragraph
                        paragraph_parent.insert(paragraph_element.getparent().index(paragraph_element) + 1, 
                                              table._element)
                        
                        # Add a paragraph after the table to prevent tables from merging
                        p = doc.add_paragraph()
                        p._element.getparent().insert(paragraph_element.getparent().index(paragraph_element) + 2, p._element)
                        
                        # Count as processed (we processed the entire keyword that produced the table)
                        processed_keywords_count += 1
                    except Exception as e:
                        st.error(f"Error inserting table: {str(e)}")
                        # If table insertion fails, keep original text
                        paragraph.text = original_text
                # Handle the old-style placeholder format for backwards compatibility
                elif isinstance(parsed_result, str) and "[TABLE_INSERTED]" in parsed_result:
                    # Check if the keyword was the only content (strip spaces for check)
                    is_only_keyword = False
                    matches = list(re.finditer(pattern, original_text))
                    if len(matches) == 1 and matches[0].group(0).strip() == original_text.strip():
                         is_only_keyword = True

                    if is_only_keyword:
                         paragraph.text = "" # Clear paragraph if only table keyword was present
                    else:
                         # Remove placeholder but keep other text
                         paragraph.text = parsed_result.replace("[TABLE_INSERTED]", "").strip()
                         
                    # Count as processed
                    processed_keywords_count += 1
                elif parsed_result != original_text:
                    paragraph.text = parsed_result
                    
                    # Estimate progress - count keywords *remaining* after parse
                    keywords_remaining = len(re.findall(pattern, paragraph.text))
                    processed_in_step = keywords_in_para - keywords_remaining
                    processed_keywords_count += processed_in_step

            except Exception as e:
                st.error(f"Error processing content '{original_text[:50]}...': {str(e)}")
                # Keep original text on error

        elements_processed += 1
        progress = elements_processed / total_elements if total_elements > 0 else 0
        progress_bar.progress(progress)

    progress_bar.progress(1.0)
    progress_text.text(f"Processing finished. Approximately {processed_keywords_count} keywords processed.")

    # Close all Excel managers to free resources
    if hasattr(parser, 'excel_managers') and parser.excel_managers:
        for filename, manager in parser.excel_managers.items():
            try:
                manager.close()
                logger.info(f"Closed Excel manager for {filename}")
            except Exception as e:
                logger.error(f"Error closing Excel manager for {filename}: {str(e)}")

    return doc, processed_keywords_count


def display_keyword_summary(summary):
    """Display analysis summary with updated Excel categories and template details."""
    st.write(f"Total keywords found: **{summary['total_keywords']}**")
    with st.expander("Document Analysis Summary"):
        col1, col2 = st.columns(2)

        with col1:
            st.markdown("**Excel Keywords (`XL!`)**")
            total_excel = sum(summary["excel_counts"].values())
            st.write(f"Total: {total_excel}")
            if summary["needs_excel"]:
                st.write("*Excel file required*")
                
                # Show Excel files if using the new format
                if "excel_files" in summary and summary["excel_files"]:
                    st.write("**Excel Files Referenced:**")
                    for excel_file in summary["excel_files"]:
                        if "excel_files_not_found" in summary and excel_file in summary["excel_files_not_found"]:
                            st.write(f"- {excel_file} (not found)")
                        else:
                            st.write(f"- {excel_file}")
                            
            # Show Excel keyword types            
            for subtype, count in summary["excel_counts"].items():
                 if count > 0: st.write(f"- {subtype}: {count}")

            st.markdown("**Input Keywords (`INPUT!`)**")
            total_inputs = sum(summary["input_counts"].values())
            st.write(f"Total: {total_inputs}")
            for input_type, count in summary["input_counts"].items():
                 if count > 0: st.write(f"- {input_type}: {count}")
                 
            st.markdown("**JSON Keywords (`JSON!`)**")
            st.write(f"Total: {summary['json_count']}")
            
            # Show JSON files
            if summary.get("needs_json") and "json_files" in summary and summary["json_files"]:
                st.write("**JSON Files Referenced:**")
                for json_file in summary["json_files"]:
                    if "json_files_not_found" in summary and json_file in summary["json_files_not_found"]:
                        st.write(f"- {json_file} (not found)")
                    else:
                        st.write(f"- {json_file}")
            
            if summary['json_count'] > 0 and 'keywords' in summary and summary['keywords']['json']:
                st.caption("Examples:")
                for item in summary['keywords']['json'][:2]:  # Show first 2
                    st.caption(f"- `{{{{{item}}}}}`")

        with col2:
            # Enhanced Template section with more details
            st.markdown("**Template Keywords (`TEMPLATE!`)**")
            st.write(f"Total: {summary['template_total']}")
            
            # Show Template files
            if summary.get("needs_templates") and "template_files" in summary and summary["template_files"]:
                st.write("**Template Files Referenced:**")
                for template_file in summary["template_files"]:
                    if "template_files_not_found" in summary and template_file in summary["template_files_not_found"]:
                        st.write(f"- {template_file} (not found)")
                    else:
                        st.write(f"- {template_file}")
            
            # Show Template keyword types
            for template_type, count in summary["template_count"].items():
                if count > 0: 
                    # Just show the uppercase name and count
                    st.write(f"- {template_type.upper()}: {count}")
            
            st.markdown("**AI Keywords (`AI!`)**")
            st.write(f"Total: {summary['ai_count']}")
            
            # Show AI files
            if summary.get("needs_ai"):
                if "ai_source_files" in summary and summary["ai_source_files"]:
                    st.write("**AI Source Files Referenced:**")
                    for ai_file in summary["ai_source_files"]:
                        if "ai_source_files_not_found" in summary and ai_file in summary["ai_source_files_not_found"]:
                            st.write(f"- {ai_file} (not found)")
                        else:
                            st.write(f"- {ai_file}")
                            
                if "ai_prompt_files" in summary and summary["ai_prompt_files"]:
                    st.write("**AI Prompt Files Referenced:**")
                    for prompt_file in summary["ai_prompt_files"]:
                        if "ai_prompt_files_not_found" in summary and prompt_file in summary["ai_prompt_files_not_found"]:
                            st.write(f"- {prompt_file} (not found)")
                        else:
                            st.write(f"- {prompt_file}")
            
            if summary['ai_count'] > 0 and 'keywords' in summary and summary['keywords']['ai']:
                st.caption("Examples:")
                for item in summary['keywords']['ai'][:2]:  # Show first 2
                    st.caption(f"- `{{{{{item}}}}}`")
            
            st.markdown("**Other/Invalid**")
            st.write(f"Total: {summary['other_count']}")
            if summary['other_count'] > 0 and 'keywords' in summary and summary['keywords']['other']:
                st.caption("Examples:")
                for item in summary['keywords']['other'][:2]:  # Show first 2
                    st.caption(f"- `{{{{{item}}}}}`")


def main():
    # Load custom CSS
    with open('style.css') as f:
        st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)
    
    logger.info("Application started")
    
    # Check for OpenAI API key at startup
    if 'api_key_checked' not in st.session_state:
        st.session_state['api_key_checked'] = False
    
    if not st.session_state['api_key_checked']:
        api_key_set = check_openai_api_key()
        
        if not api_key_set:
            st.warning("OpenAI API key is not set. Please enter your API key below.")
            st.info("For security, your key will only be stored in this session and not saved to disk.")
            api_key = st.text_input("OpenAI API Key", type="password", help="Your key will only be stored in memory for this session")
            
            if st.button("Use API Key"):
                if api_key:
                    # Store API key in session state only
                    st.session_state['openai_api_key'] = api_key
                    st.session_state['api_key_checked'] = True
                    st.success("API key set for this session!")
                    st.rerun()
                else:
                    st.error("Please enter a valid API key.")
            
            # Stop further execution until API key is provided
            st.info("Please provide an OpenAI API key to continue.")
            return
        
        st.session_state['api_key_checked'] = True
    
    # Initialize parser instance for help text display
    if 'keyword_parser_instance_for_help' not in st.session_state:
        st.session_state.keyword_parser_instance_for_help = keywordParser()
    
    # --- State Management ---
    # Initialize state variables if they don't exist
    default_state = {
        'current_step': 1,  # Track the current wizard step
        'doc_uploaded': False, 'doc_path': None, 'analysis_summary': None,
        'excel_uploaded': False, 'excel_path': None, 'excel_manager_instance': None,
        'excel_files_uploaded': {}, 'excel_managers': {},  # New state for multiple Excel files
        'templates_uploaded': False, 'template_files_uploaded': {}, # New state for templates
        'json_uploaded': False, 'json_files_uploaded': {}, # New state for JSON files
        'ai_uploaded': False, 'ai_source_files_uploaded': {}, 'ai_prompt_files_uploaded': {}, # New state for AI files
        'rerun_triggered_after_upload': False, 'rerun_triggered_for_found_files': False,  # Flags to prevent infinite reruns
        'rerun_triggered_after_template_upload': False, 'rerun_triggered_for_found_templates': False,  # Template flags
        'rerun_triggered_after_json_upload': False, 'rerun_triggered_for_found_json': False,  # JSON flags
        'rerun_triggered_after_ai_upload': False, 'rerun_triggered_for_found_ai': False,  # AI flags
        'keyword_parser_instance': None, 'form_submitted_main': False, 'input_values_main': {},
        'processing_started': False, 'processed_doc_path': None, 'processed_count': 0
    }
    for key, value in default_state.items():
        if key not in st.session_state:
            st.session_state[key] = value
    
    # Sidebar with keyword reference guide and reset button
    with st.sidebar:
        # Load and display the Form Filler logo
        st.image("assets/images/form_filler_logo.png", width=250)
        
        st.subheader("Navigation")
        
        # Add visual step indicator
        st.markdown("""
        <div class="step-indicator">
            <div class="step {0}">1</div>
            <div class="step-line"></div>
            <div class="step {1}">2</div>
            <div class="step-line"></div>
            <div class="step {2}">3</div>
            <div class="step-line"></div>
            <div class="step {3}">4</div>
            <div class="step-line"></div>
            <div class="step {4}">5</div>
        </div>
        """.format(
            "active" if st.session_state.current_step == 1 else "",
            "active" if st.session_state.current_step == 2 else "",
            "active" if st.session_state.current_step == 3 else "",
            "active" if st.session_state.current_step == 4 else "",
            "active" if st.session_state.current_step == 5 else ""
        ), unsafe_allow_html=True)
        
        # Step indicator
        st.write("Current step: ", st.session_state.current_step)
        
        # Add Navigation buttons
        if st.session_state.current_step > 1:
            if st.button("← Previous Step"):
                st.session_state.current_step -= 1
                st.rerun()
        
        # Only show Next button if it makes sense for the current step
        can_proceed = False
        if st.session_state.current_step == 1 and st.session_state.doc_uploaded:
            can_proceed = True
        elif st.session_state.current_step == 2:
            needs_excel = st.session_state.analysis_summary and st.session_state.analysis_summary.get("needs_excel", False)
            needs_templates = st.session_state.analysis_summary and st.session_state.analysis_summary.get("needs_templates", False)
            needs_json = st.session_state.analysis_summary and st.session_state.analysis_summary.get("needs_json", False)
            needs_ai = st.session_state.analysis_summary and st.session_state.analysis_summary.get("needs_ai", False)
            
            excel_ready = (not needs_excel) or st.session_state.excel_uploaded
            templates_ready = (not needs_templates) or st.session_state.templates_uploaded
            json_ready = (not needs_json) or st.session_state.json_uploaded
            ai_ready = (not needs_ai) or st.session_state.ai_uploaded
            
            can_proceed = excel_ready and templates_ready and json_ready and ai_ready
        elif st.session_state.current_step == 3:
            has_inputs = st.session_state.analysis_summary and sum(st.session_state.analysis_summary['input_counts'].values()) > 0
            can_proceed = (not has_inputs) or st.session_state.form_submitted_main
        
        if can_proceed and st.session_state.current_step < 5:
            if st.button("Next Step →"):
                st.session_state.current_step += 1
                st.rerun()

        st.subheader("Keyword Reference Guides")
        # Keyword reference guide in expandable section
        with st.expander("Input Keyword Reference Guide", expanded=False):
            st.markdown(st.session_state.keyword_parser_instance_for_help.get_input_keyword_help())
        with st.expander("Excel Keyword Reference Guide", expanded=False):
            st.markdown(st.session_state.keyword_parser_instance_for_help.get_excel_keyword_help())
        with st.expander("Template Keyword Reference Guide", expanded=False):
            st.markdown(st.session_state.keyword_parser_instance_for_help.get_template_keyword_help())
        with st.expander("JSON Keyword Reference Guide", expanded=False):
            st.markdown(st.session_state.keyword_parser_instance_for_help.get_json_keyword_help())
        with st.expander("AI Keyword Reference Guide", expanded=False):
            st.markdown(st.session_state.keyword_parser_instance_for_help.get_ai_keyword_help())

        # Reset button
        if st.button("Reset Application"):
            logger.info("Resetting application state")
            # Clean up temp files
            if st.session_state.doc_path and os.path.exists(st.session_state.doc_path): 
                os.unlink(st.session_state.doc_path)
                logger.info(f"Removed temporary document: {st.session_state.doc_path}")
            if st.session_state.excel_path and os.path.exists(st.session_state.excel_path): 
                os.unlink(st.session_state.excel_path)
                logger.info(f"Removed temporary Excel file: {st.session_state.excel_path}")
            if st.session_state.processed_doc_path and os.path.exists(st.session_state.processed_doc_path): 
                os.unlink(st.session_state.processed_doc_path)
                logger.info(f"Removed processed document: {st.session_state.processed_doc_path}")
            
            # Close Excel Managers
            if st.session_state.excel_manager_instance: 
                st.session_state.excel_manager_instance.close()
                
            # Close all individual Excel managers
            if 'excel_managers' in st.session_state and st.session_state.excel_managers:
                for filename, manager in st.session_state.excel_managers.items():
                    try:
                        manager.close()
                        logger.info(f"Closed Excel manager for {filename}")
                    except Exception as e:
                        logger.error(f"Error closing Excel manager for {filename}: {str(e)}")
            
            # Reset state variables
            for key in default_state:
                st.session_state[key] = default_state[key]
                
            # Clear additional Excel-related state
            if 'excel_files_uploaded' in st.session_state:
                st.session_state.excel_files_uploaded = {}
            if 'excel_managers' in st.session_state:
                st.session_state.excel_managers = {}
                
            # Clear template-related state
            if 'template_files_uploaded' in st.session_state:
                st.session_state.template_files_uploaded = {}
            st.session_state.templates_uploaded = False
                
            # Clear JSON-related state
            if 'json_files_uploaded' in st.session_state:
                st.session_state.json_files_uploaded = {}
            st.session_state.json_uploaded = False
                
            # Clear AI-related state
            if 'ai_source_files_uploaded' in st.session_state:
                st.session_state.ai_source_files_uploaded = {}
            if 'ai_prompt_files_uploaded' in st.session_state:
                st.session_state.ai_prompt_files_uploaded = {}
            st.session_state.ai_uploaded = False
                
            # Reset rerun flags
            st.session_state.rerun_triggered_after_upload = False
            st.session_state.rerun_triggered_for_found_files = False
            st.session_state.rerun_triggered_after_template_upload = False
            st.session_state.rerun_triggered_for_found_templates = False
            st.session_state.rerun_triggered_after_json_upload = False
            st.session_state.rerun_triggered_for_found_json = False
            st.session_state.rerun_triggered_after_ai_upload = False
            st.session_state.rerun_triggered_for_found_ai = False
                
            st.rerun()
    
    # Main content area - only show the current step
    # --- Step 1: Upload Word Document ---
    if st.session_state.current_step == 1:
        st.header("Step 1: Upload Document")
        st.write("Upload a Word document containing keywords that you want to process. The document should include keywords in double curly braces like `{{keyword}}`. Refer to the Keyword Reference Guides in the sidebar for Keyword Help.")
        
        doc_file = st.file_uploader("Upload Word Document (.docx)", type=["docx"], key="main_doc_uploader")

        if doc_file and not st.session_state.doc_uploaded:
            # Reset relevant states for new upload
            st.session_state.update({k: v for k, v in default_state.items() 
                                   if k not in ['keyword_parser_instance_for_help', 'current_step']})
            st.session_state.current_step = 1  # Stay on step 1
            
            # Save new doc
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_doc:
                tmp_doc.write(doc_file.getvalue())
                st.session_state.doc_path = tmp_doc.name
            st.session_state.doc_uploaded = True
            st.rerun()
        
    
    # --- Step 2: Analysis & Excel Upload (if needed) ---
    elif st.session_state.current_step == 2:
        st.header("Step 2: Document Analysis & Required Files")
        st.write("This step analyzes your document to identify keywords and determines if additional files (like Excel, Templates, JSON, AI Source, or AI Prompt files) are needed.")
        
        # First run analysis if needed
        if not st.session_state.analysis_summary:
            with st.spinner("Analyzing document..."):
                try:
                    summary = preprocess_word_doc(st.session_state.doc_path)
                    st.session_state.analysis_summary = summary
                    
                    # Initialize session state for excel files
                    if "excel_files_uploaded" not in st.session_state:
                        st.session_state.excel_files_uploaded = {}
                    if "excel_managers" not in st.session_state:
                        st.session_state.excel_managers = {}
                    
                    st.rerun()
                except Exception as e:
                    st.error(f"Analysis failed: {e}")
                    st.session_state.doc_uploaded = False  # Allow re-upload
                    st.session_state.current_step = 1  # Go back to step 1
                    st.rerun()
        
        # Display analysis results
        if st.session_state.analysis_summary:
            display_keyword_summary(st.session_state.analysis_summary)
            needs_excel = st.session_state.analysis_summary["needs_excel"]
            
            # Only show Excel uploader if needed based on analysis
            if needs_excel:
                if st.session_state.analysis_summary.get("excel_files"):
                    # New format with specific Excel files
                    excel_files = st.session_state.analysis_summary["excel_files"]
                    excel_files_not_found = st.session_state.analysis_summary["excel_files_not_found"]
                    
                    if excel_files_not_found:
                        st.write("### Excel File(s) Required")
                        st.write("The following Excel file(s) were specified in the document but not found. Please upload them:")
                        
                        for excel_file in excel_files_not_found:
                            # Check if this file has already been uploaded
                            if excel_file in st.session_state.excel_files_uploaded:
                                st.success(f"✅ {excel_file} has been uploaded.")
                                continue
                                
                            st.write(f"**{excel_file}**")
                            excel_upload_key = f"excel_uploader_{excel_file}"
                            
                            # Create uploader for this file
                            uploaded_file = st.file_uploader(f"Upload {excel_file}", 
                                                           type=["xlsx", "xls"], key=excel_upload_key)
                            
                            if uploaded_file:
                                # Save the uploaded file to the excel directory with the exact filename specified
                                excel_dir = "excel"
                                save_path = os.path.join(excel_dir, excel_file)
                                
                                with open(save_path, "wb") as f:
                                    f.write(uploaded_file.getbuffer())
                                
                                st.success(f"Saved {excel_file} to excel folder")
                                logger.info(f"Saved uploaded Excel file to {save_path}")
                                
                                # Initialize Excel manager for this file
                                try:
                                    st.session_state.excel_managers[excel_file] = excelManager(save_path)
                                    st.session_state.excel_files_uploaded[excel_file] = True
                                    st.rerun()  # Refresh to update the UI
                                except Exception as e:
                                    st.error(f"Failed to load Excel file {excel_file}: {e}")
                                    logger.error(f"Failed to load Excel file {excel_file}: {e}", exc_info=True)
                        
                        # Check if all required files have been uploaded
                        all_files_uploaded = all(excel_file in st.session_state.excel_files_uploaded 
                                              for excel_file in excel_files_not_found)
                        
                        if all_files_uploaded:
                            st.success("All required Excel files have been uploaded!")
                            st.session_state.excel_uploaded = True
                            # Only rerun if this is the first time we're setting the flag
                            if not st.session_state.get('rerun_triggered_after_upload', False):
                                st.session_state.rerun_triggered_after_upload = True
                                st.rerun()
                        else:
                            st.session_state.excel_uploaded = False
                            
                    else:
                        # All specified Excel files were found
                        st.success("All Excel files specified in the document have been found in the excel folder.")
                        st.session_state.excel_uploaded = True
                        
                        # Initialize managers for the found files
                        excel_dir = "excel"
                        for excel_file in excel_files:
                            file_path = os.path.join(excel_dir, excel_file)
                            if os.path.exists(file_path) and excel_file not in st.session_state.excel_managers:
                                try:
                                    st.session_state.excel_managers[excel_file] = excelManager(file_path)
                                    st.session_state.excel_files_uploaded[excel_file] = True
                                except Exception as e:
                                    st.error(f"Failed to load Excel file {excel_file}: {e}")
                                    logger.error(f"Failed to load Excel file {excel_file}: {e}", exc_info=True)
                        
                        # Only rerun if this is the first time we're setting the flag for found files
                        if not st.session_state.get('rerun_triggered_for_found_files', False):
                            st.session_state.rerun_triggered_for_found_files = True
                            st.rerun()
                else:
                    # Old format without specific Excel files - use the standard uploader
                    st.write("Based on the analysis, an Excel file is required.")
                    st.write("*Note: The document uses the old Excel keyword format without specifying files. Please upload an Excel file to use for all Excel keywords.*")
                    
                    excel_file = st.file_uploader("Upload Required Excel Spreadsheet (.xlsx)", 
                                                type=["xlsx"], key="main_excel_uploader")
                    
                    if excel_file and not st.session_state.excel_uploaded:
                        # Save new excel file
                        if st.session_state.excel_path and os.path.exists(st.session_state.excel_path): 
                            os.unlink(st.session_state.excel_path)  # Clean old temp excel
                        
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_excel:
                            tmp_excel.write(excel_file.getvalue())
                            st.session_state.excel_path = tmp_excel.name
                        
                        st.session_state.excel_uploaded = True
                        
                        # Reset excel manager instance as file changed
                        if st.session_state.excel_manager_instance: 
                            st.session_state.excel_manager_instance.close()
                        st.session_state.excel_manager_instance = None
                        st.rerun()
            else:
                st.success("No Excel file required. You can proceed to the next step.")
                
            # Check if template files are needed
            needs_templates = st.session_state.analysis_summary.get("needs_templates", False)
            
            if needs_templates and "template_files" in st.session_state.analysis_summary:
                template_files = st.session_state.analysis_summary["template_files"]
                template_files_not_found = st.session_state.analysis_summary.get("template_files_not_found", [])
                
                if template_files_not_found:
                    st.write("### Template File(s) Required")
                    st.write("The following template file(s) were specified in the document but not found. Please upload them:")
                    
                    for template_file in template_files_not_found:
                        # Check if this file has already been uploaded
                        if template_file in st.session_state.template_files_uploaded:
                            st.success(f"✅ {template_file} has been uploaded.")
                            continue
                            
                        st.write(f"**{template_file}**")
                        template_upload_key = f"template_uploader_{template_file}"
                        
                        # Create uploader for this file
                        uploaded_file = st.file_uploader(f"Upload {template_file}", 
                                                       type=["docx", "txt"], key=template_upload_key)
                        
                        if uploaded_file:
                            # Save the uploaded file to the templates directory with the exact filename specified
                            templates_dir = "templates"
                            save_path = os.path.join(templates_dir, template_file)
                            
                            with open(save_path, "wb") as f:
                                f.write(uploaded_file.getbuffer())
                            
                            st.success(f"Saved {template_file} to templates folder")
                            logger.info(f"Saved uploaded template file to {save_path}")
                            
                            # Mark template as uploaded
                            st.session_state.template_files_uploaded[template_file] = True
                            st.rerun()  # Refresh to update the UI
                    
                    # Check if all required template files have been uploaded
                    all_templates_uploaded = all(template_file in st.session_state.template_files_uploaded 
                                          for template_file in template_files_not_found)
                    
                    if all_templates_uploaded:
                        st.success("All required template files have been uploaded!")
                        st.session_state.templates_uploaded = True
                        # Only rerun if this is the first time we're setting the flag
                        if not st.session_state.get('rerun_triggered_after_template_upload', False):
                            st.session_state.rerun_triggered_after_template_upload = True
                            st.rerun()
                    else:
                        st.session_state.templates_uploaded = False
                else:
                    # All template files were found
                    st.success("All template files specified in the document have been found in the templates folder.")
                    st.session_state.templates_uploaded = True
                    
                    # Only rerun if this is the first time we're setting the flag for found templates
                    if not st.session_state.get('rerun_triggered_for_found_templates', False):
                        st.session_state.rerun_triggered_for_found_templates = True
                        st.rerun()
            
            # Check if JSON files are needed
            needs_json = st.session_state.analysis_summary.get("needs_json", False)
            
            if needs_json and "json_files" in st.session_state.analysis_summary:
                json_files = st.session_state.analysis_summary["json_files"]
                json_files_not_found = st.session_state.analysis_summary.get("json_files_not_found", [])
                
                if json_files_not_found:
                    st.write("### JSON File(s) Required")
                    st.write("The following JSON file(s) were specified in the document but not found. Please upload them:")
                    
                    for json_file in json_files_not_found:
                        # Check if this file has already been uploaded
                        if json_file in st.session_state.json_files_uploaded:
                            st.success(f"✅ {json_file} has been uploaded.")
                            continue
                            
                        st.write(f"**{json_file}**")
                        json_upload_key = f"json_uploader_{json_file}"
                        
                        # Create uploader for this file
                        uploaded_file = st.file_uploader(f"Upload {json_file}", 
                                                       type=["json"], key=json_upload_key)
                        
                        if uploaded_file:
                            # Save the uploaded file to the json directory with the exact filename specified
                            json_dir = "json"
                            save_path = os.path.join(json_dir, json_file)
                            
                            with open(save_path, "wb") as f:
                                f.write(uploaded_file.getbuffer())
                            
                            st.success(f"Saved {json_file} to json folder")
                            logger.info(f"Saved uploaded JSON file to {save_path}")
                            
                            # Mark JSON as uploaded
                            st.session_state.json_files_uploaded[json_file] = True
                            st.rerun()  # Refresh to update the UI
                    
                    # Check if all required JSON files have been uploaded
                    all_json_uploaded = all(json_file in st.session_state.json_files_uploaded 
                                           for json_file in json_files_not_found)
                    
                    if all_json_uploaded:
                        st.success("All required JSON files have been uploaded!")
                        st.session_state.json_uploaded = True
                        # Only rerun if this is the first time we're setting the flag
                        if not st.session_state.get('rerun_triggered_after_json_upload', False):
                            st.session_state.rerun_triggered_after_json_upload = True
                            st.rerun()
                    else:
                        st.session_state.json_uploaded = False
                else:
                    # All JSON files were found
                    st.success("All JSON files specified in the document have been found in the json folder.")
                    st.session_state.json_uploaded = True
                    
                    # Only rerun if this is the first time we're setting the flag for found JSON
                    if not st.session_state.get('rerun_triggered_for_found_json', False):
                        st.session_state.rerun_triggered_for_found_json = True
                        st.rerun()
            
            # Check if AI files are needed
            needs_ai = st.session_state.analysis_summary.get("needs_ai", False)
            
            if needs_ai:
                # Handle AI source files
                ai_source_files = st.session_state.analysis_summary.get("ai_source_files", [])
                ai_source_files_not_found = st.session_state.analysis_summary.get("ai_source_files_not_found", [])
                
                if ai_source_files_not_found:
                    st.write("### AI Source File(s) Required")
                    st.write("The following AI source file(s) were specified in the document but not found. Please upload them:")
                    
                    for ai_file in ai_source_files_not_found:
                        # Check if this file has already been uploaded
                        if ai_file in st.session_state.ai_source_files_uploaded:
                            st.success(f"✅ {ai_file} has been uploaded.")
                            continue
                            
                        st.write(f"**{ai_file}**")
                        ai_upload_key = f"ai_source_uploader_{ai_file}"
                        
                        # Determine file type based on extension
                        file_ext = ai_file.split('.')[-1].lower() if '.' in ai_file else 'docx'
                        allowed_types = ["docx", "txt"] if file_ext in ["docx", "txt"] else ["docx", "txt", file_ext]
                        
                        # Create uploader for this file
                        uploaded_file = st.file_uploader(f"Upload {ai_file}", 
                                                       type=allowed_types, key=ai_upload_key)
                        
                        if uploaded_file:
                            # Save the uploaded file to the ai directory with the exact filename specified
                            ai_dir = "ai"
                            save_path = os.path.join(ai_dir, ai_file)
                            
                            with open(save_path, "wb") as f:
                                f.write(uploaded_file.getbuffer())
                            
                            st.success(f"Saved {ai_file} to ai folder")
                            logger.info(f"Saved uploaded AI source file to {save_path}")
                            
                            # Mark AI source file as uploaded
                            st.session_state.ai_source_files_uploaded[ai_file] = True
                            st.rerun()  # Refresh to update the UI
                
                # Handle AI prompt files
                ai_prompt_files = st.session_state.analysis_summary.get("ai_prompt_files", [])
                ai_prompt_files_not_found = st.session_state.analysis_summary.get("ai_prompt_files_not_found", [])
                
                if ai_prompt_files_not_found:
                    st.write("### AI Prompt File(s) Required")
                    st.write("The following AI prompt file(s) were specified in the document but not found. Please upload them:")
                    
                    for prompt_file in ai_prompt_files_not_found:
                        # Check if this file has already been uploaded
                        if prompt_file in st.session_state.ai_prompt_files_uploaded:
                            st.success(f"✅ {prompt_file} has been uploaded.")
                            continue
                            
                        st.write(f"**{prompt_file}**")
                        prompt_upload_key = f"ai_prompt_uploader_{prompt_file}"
                        
                        # Create uploader for this file
                        uploaded_file = st.file_uploader(f"Upload {prompt_file}", 
                                                       type=["txt"], key=prompt_upload_key)
                        
                        if uploaded_file:
                            # Save the uploaded file to the ai directory with the exact filename specified
                            ai_dir = "ai"
                            save_path = os.path.join(ai_dir, prompt_file)
                            
                            with open(save_path, "wb") as f:
                                f.write(uploaded_file.getbuffer())
                            
                            st.success(f"Saved {prompt_file} to ai folder")
                            logger.info(f"Saved uploaded AI prompt file to {save_path}")
                            
                            # Mark AI prompt file as uploaded
                            st.session_state.ai_prompt_files_uploaded[prompt_file] = True
                            st.rerun()  # Refresh to update the UI
                
                # Check if all required AI files have been uploaded
                all_ai_source_uploaded = all(ai_file in st.session_state.ai_source_files_uploaded 
                                          for ai_file in ai_source_files_not_found)
                all_ai_prompt_uploaded = all(prompt_file in st.session_state.ai_prompt_files_uploaded 
                                          for prompt_file in ai_prompt_files_not_found)
                
                if all_ai_source_uploaded and all_ai_prompt_uploaded:
                    st.success("All required AI files have been uploaded!")
                    st.session_state.ai_uploaded = True
                    # Only rerun if this is the first time we're setting the flag
                    if not st.session_state.get('rerun_triggered_after_ai_upload', False):
                        st.session_state.rerun_triggered_after_ai_upload = True
                        st.rerun()
                else:
                    st.session_state.ai_uploaded = False
                    # Show status of what's still missing
                    if ai_source_files_not_found and not all_ai_source_uploaded:
                        missing_source = [f for f in ai_source_files_not_found if f not in st.session_state.ai_source_files_uploaded]
                        #st.info(f"Still waiting for AI source files: {', '.join(missing_source)}")
                    if ai_prompt_files_not_found and not all_ai_prompt_uploaded:
                        missing_prompts = [f for f in ai_prompt_files_not_found if f not in st.session_state.ai_prompt_files_uploaded]
                        #st.info(f"Still waiting for AI prompt files: {', '.join(missing_prompts)}")
                
                # If no missing files were found but we need AI
                if not ai_source_files_not_found and not ai_prompt_files_not_found and ai_source_files:
                    st.success("All AI files specified in the document have been found in the ai folder.")
                    st.session_state.ai_uploaded = True
                    
                    # Only rerun if this is the first time we're setting the flag for found AI
                    if not st.session_state.get('rerun_triggered_for_found_ai', False):
                        st.session_state.rerun_triggered_for_found_ai = True
                        st.rerun()
            
            # Initialize Excel Manager for old format
            if needs_excel and not st.session_state.analysis_summary.get("excel_files") and st.session_state.excel_path and not st.session_state.excel_manager_instance:
                try:
                    with st.spinner("Loading Excel data..."):
                        st.session_state.excel_manager_instance = excelManager(st.session_state.excel_path)
                except Exception as e:
                    st.error(f"Failed to load Excel file: {e}")
                    st.session_state.excel_uploaded = False  # Reset upload status
                    st.rerun()
            
            # Create parser instance with appropriate Excel manager(s)
            # This will be a more complex parser setup for the new format to handle multiple Excel files
            current_excel_manager = None
            if needs_excel:
                if st.session_state.analysis_summary.get("excel_files"):
                    # New format - use the first manager as default but will update parser to handle all files
                    if st.session_state.excel_managers:
                        # Use the first manager as the default
                        current_excel_manager = next(iter(st.session_state.excel_managers.values()))
                else:
                    # Old format
                    current_excel_manager = st.session_state.excel_manager_instance
            
            # Always ensure parser instance exists, update if excel manager changes
            if not st.session_state.keyword_parser_instance or getattr(st.session_state.keyword_parser_instance, 'excel_manager', None) != current_excel_manager:
                st.session_state.keyword_parser_instance = keywordParser(current_excel_manager)
                
                # If we have multiple Excel managers, store them for access in the parser
                if st.session_state.excel_managers:
                    st.session_state.keyword_parser_instance.excel_managers = st.session_state.excel_managers
    
    # --- Step 3: User Input Form (if needed) ---
    elif st.session_state.current_step == 3:
        st.header("Step 3: Provide Input Values")
        st.write("Fill in values for the input fields found in your document. These values will replace the corresponding keywords during processing.")
        
        # Check if inputs are needed
        has_inputs = st.session_state.analysis_summary and sum(st.session_state.analysis_summary['input_counts'].values()) > 0
        
        if not has_inputs:
            st.success("No user inputs required.")
            #st.info("Click 'Next Step →' in the sidebar to continue to processing.")
        else:
            st.write("Please provide values for the keywords found in your document.")
            
            if not st.session_state.form_submitted_main:
                parser = st.session_state.keyword_parser_instance
                
                with st.form(key="main_input_form"):
                    # Use analysis summary to find all input keywords
                    all_input_keywords = [item for sublist in st.session_state.analysis_summary['keywords']['input'].values() 
                                         for item in sublist]
                    unique_input_contents = sorted(list(set(all_input_keywords)))  # Get unique input definitions
                    
                    # Store fields in local state for this form
                    temp_inputs = {}
                    
                    # Disable the parser's internal form handling to prevent duplication
                    parser.form_submitted = True
                    
                    for content in unique_input_contents:
                        # Create field using parser's helper function
                        field_key = f"input_field_{content}"
                        temp_inputs[content] = parser._create_input_field(content)
                    
                    submitted = st.form_submit_button("Submit Inputs")
                    if submitted:
                        # Store values in session state
                        for content in unique_input_contents:
                            field_key = f"input_field_{content}"
                            # Extract field values from session state
                            if field_key in st.session_state:
                                field_value = st.session_state[field_key]
                                # Ensure we're storing the content exactly as it appears in the document
                                st.session_state.input_values_main[content] = field_value
                        
                        # Update the parser's internal values - ensure we use the full keyword format
                        for content, value in st.session_state.input_values_main.items():
                            # Store using the full format with INPUT!
                            keyword = f"{{{{{content}}}}}"
                            parser.input_values[keyword] = value
                            
                            # Also store in alternate formats to maximize chances of matching
                            if content.startswith("INPUT!"):
                                # Also store without the INPUT! prefix
                                non_prefix_content = content[6:]  # Remove "INPUT!"
                                alt_keyword = f"{{{{{non_prefix_content}}}}}"
                                parser.input_values[alt_keyword] = value
                            else:
                                # Also store with the INPUT! prefix
                                alt_keyword = f"{{{{INPUT!{content}}}}}"
                                parser.input_values[alt_keyword] = value
                        
                        st.session_state.form_submitted_main = True
                        logger.info("Form inputs submitted")
                        st.rerun()
            else:
                st.success("Input values submitted successfully!")
                #st.info("Click 'Next Step →' in the sidebar to continue to processing.")
    
    # --- Step 4: Process Document ---
    elif st.session_state.current_step == 4:
        st.header("Step 4: Process Document")
        st.write("Now the system will replace all keywords in your document with their corresponding values from User Inputs, Excel, Templates, JSON, and AI Keywords.")
        
        # Determine if ready to process
        needs_excel = st.session_state.analysis_summary and st.session_state.analysis_summary["needs_excel"]
        has_inputs = st.session_state.analysis_summary and sum(st.session_state.analysis_summary['input_counts'].values()) > 0
        
        ready_to_process = st.session_state.doc_uploaded and \
                          (not needs_excel or st.session_state.excel_uploaded) and \
                          (not has_inputs or st.session_state.form_submitted_main)
        
        process_button_disabled = not ready_to_process or st.session_state.processing_started
        
        if not process_button_disabled:
            if st.button("Process Document Now", key="main_process_btn"):
                st.session_state.processing_started = True
                st.session_state.processed_doc_path = None  # Clear previous
                
                with st.spinner("Processing document... This may take a moment."):
                    try:
                        # Ensure parser has the submitted inputs
                        parser = st.session_state.keyword_parser_instance
                        
                        # Force the parser to use our input values, not its internal form
                        parser.form_submitted = True
                        
                        # Process the document
                        progress_bar = st.progress(0)
                        progress_text = st.empty()
                        progress_text.text("Processing keywords...")
                        
                        processed_doc, count = process_word_doc(
                            st.session_state.doc_path,
                            st.session_state.excel_path,
                            parser=parser
                        )
                        
                        if processed_doc:
                            tmp_folder = "tmp"
                            if not os.path.exists(tmp_folder): os.makedirs(tmp_folder)
                            # Use original filename for output
                            base_name = os.path.basename(st.session_state.doc_path)
                            output_filename = f"processed_{base_name}" if not base_name.startswith("tmp") else "processed_document.docx"
                            output_path = os.path.join(tmp_folder, output_filename)
                            
                            processed_doc.save(output_path)
                            st.session_state.processed_doc_path = output_path
                            st.session_state.processed_count = count
                            st.success(f"Processing Complete! Approximately {count} keywords processed.")
                            logger.info(f"Document processed successfully. Saved to {output_path}. {count} keywords processed.")
                            
                            # Automatically move to the download step
                            st.session_state.current_step = 5
                            st.rerun()
                        else:
                            st.warning("Processing did not return a document.")
                    
                    except Exception as e:
                        st.error(f"Error during processing: {e}")
                        logger.error(f"Processing error: {str(e)}", exc_info=True)
                    finally:
                        # Close excel manager instance if it exists
                        if st.session_state.excel_manager_instance:
                            st.session_state.excel_manager_instance.close()
                            st.session_state.excel_manager_instance = None
                        st.session_state.processing_started = False  # Reset processing flag
        else:
            if not ready_to_process:
                st.warning("Please complete the previous steps before processing.")
            elif st.session_state.processing_started:
                st.info("Processing is currently in progress...")
    
    # --- Step 5: Download ---
    elif st.session_state.current_step == 5:
        st.header("Step 5: Download Result")
        st.write("Your document has been processed successfully! You can now download the final document with all keywords replaced.")
        
        if st.session_state.processed_doc_path:
            st.success(f"Document processed successfully! {st.session_state.processed_count} keywords were replaced.")
            st.write("Your document is ready to download.")
            
            try:
                with open(st.session_state.processed_doc_path, "rb") as fp:
                    st.download_button(
                        label="📥 Download Processed Document",
                        data=fp,
                        file_name=os.path.basename(st.session_state.processed_doc_path),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    st.write("You can also:")
                    if st.button("Start Over with a New Document"):
                        # Reset to initial state but keep the help parser
                        if st.session_state.doc_path and os.path.exists(st.session_state.doc_path): 
                            os.unlink(st.session_state.doc_path)
                        if st.session_state.excel_path and os.path.exists(st.session_state.excel_path): 
                            os.unlink(st.session_state.excel_path)
                        if st.session_state.processed_doc_path and os.path.exists(st.session_state.processed_doc_path):
                            os.unlink(st.session_state.processed_doc_path)
                        # Reset state except keyword_parser_instance_for_help
                        parser_for_help = st.session_state.keyword_parser_instance_for_help
                        for key in default_state:
                            st.session_state[key] = default_state[key]
                        st.session_state.keyword_parser_instance_for_help = parser_for_help
                        st.rerun()
            except FileNotFoundError:
                st.error("Processed file not found. Please try processing again.")
                st.session_state.processed_doc_path = None  # Reset path
                st.session_state.current_step = 4  # Go back to processing step
                st.rerun()
        else:
            st.error("No processed document available. Please go back to the processing step.")
            if st.button("Return to Processing Step"):
                st.session_state.current_step = 4
                st.rerun()


if __name__ == "__main__":
    main()